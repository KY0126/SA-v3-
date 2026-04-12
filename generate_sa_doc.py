#!/usr/bin/env python3
"""
Generate FJU Smart Hub SA Specification Word Document
with complete 21-table database schema and auto-generated TOC
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# ─── Color constants ───
FJU_NAVY = RGBColor(0x00, 0x31, 0x53)
FJU_GOLD = RGBColor(0xC4, 0x9A, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "003153"
SUBHEADER_BG = "D4E6F1"
ROW_ALT_BG = "F8F9FA"
FK_COLOR = RGBColor(0x1A, 0x5C, 0x8A)
ENUM_COLOR = RGBColor(0x8B, 0x00, 0x00)


def set_cell_shading(cell, color_hex):
    """Set background color of a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_table_header_row(table, headers, widths=None):
    """Format table header row with navy background."""
    row = table.rows[0]
    for i, cell in enumerate(row.cells):
        set_cell_shading(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(headers[i])
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Microsoft JhengHei'


def add_toc(doc):
    """Add an auto-generated Table of Contents field."""
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Heading 1']
    run = paragraph.add_run('目錄')

    # Add TOC field
    paragraph2 = doc.add_paragraph()
    run2 = paragraph2.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2._r.append(fldChar1)

    run3 = paragraph2.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run3._r.append(instrText)

    run4 = paragraph2.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run4._r.append(fldChar2)

    run5 = paragraph2.add_run('（請在 Word 中按右鍵 > 更新功能變數 以自動產生完整目錄）')
    run5.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run5.font.italic = True

    run6 = paragraph2.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run6._r.append(fldChar3)


def make_db_detail_table(doc, table_num, table_name, cn_name, description, columns, notes=None):
    """Create a detailed schema table for one DB table."""
    # Heading 3 for each table
    h = doc.add_heading(f'3.4.{table_num} {table_name}（{cn_name}）', level=3)

    # Description
    p = doc.add_paragraph(description)
    p.paragraph_format.space_after = Pt(6)

    # Column detail table: 欄位名, 型別, 預設值, Nullable, 說明
    headers = ['欄位名稱', '資料型別', '預設值', '可為空', '說明']
    tbl = doc.add_table(rows=1 + len(columns), cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    col_widths = [Cm(3.5), Cm(4.0), Cm(2.5), Cm(1.5), Cm(5.5)]
    for i, width in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = width

    # Header row
    for i, h_text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h_text
        set_cell_shading(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(8.5)
            run.font.name = 'Microsoft JhengHei'

    # Data rows
    for row_idx, col_data in enumerate(columns):
        row = tbl.rows[row_idx + 1]
        # Alternate row color
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, ROW_ALT_BG)

        for col_idx, value in enumerate(col_data):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            p.clear()

            # Format column name (bold for PK)
            if col_idx == 0:
                run = p.add_run(str(value))
                run.font.size = Pt(8)
                run.font.name = 'Consolas'
                if value == 'id':
                    run.font.bold = True
                    run.font.color.rgb = FJU_GOLD
            elif col_idx == 1:
                # Type column - color ENUM/FK differently
                run = p.add_run(str(value))
                run.font.size = Pt(8)
                run.font.name = 'Consolas'
                if 'ENUM' in str(value).upper():
                    run.font.color.rgb = ENUM_COLOR
                elif 'FK' in str(value).upper():
                    run.font.color.rgb = FK_COLOR
            else:
                run = p.add_run(str(value))
                run.font.size = Pt(8)
                run.font.name = 'Microsoft JhengHei'

            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    # Notes (indexes, constraints, FKs)
    if notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        for note_label, note_text in notes:
            run = p.add_run(f'📌 {note_label}：')
            run.font.bold = True
            run.font.size = Pt(8.5)
            run.font.name = 'Microsoft JhengHei'
            run = p.add_run(f'{note_text}\n')
            run.font.size = Pt(8.5)
            run.font.name = 'Microsoft JhengHei'

    doc.add_paragraph()  # spacing


def build_document():
    """Build the complete SA specification document."""
    # Load existing document to preserve styles/formatting
    src = Document('FJU_Smart_Hub_SA_Specification.docx')
    doc = Document()

    # ─── Copy styles from source ───
    # We'll create fresh but use similar formatting

    # ─── Set default font ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10.5)

    # Set heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Microsoft JhengHei'
        h_style.font.color.rgb = FJU_NAVY
        if level == 1:
            h_style.font.size = Pt(18)
        elif level == 2:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)

    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ══════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('天主教輔仁大學')
    run.font.size = Pt(22)
    run.font.color.rgb = FJU_NAVY
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('學務處 課外活動指導組')
    run.font.size = Pt(16)
    run.font.color.rgb = FJU_NAVY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FJU Smart Hub')
    run.font.size = Pt(28)
    run.font.color.rgb = FJU_GOLD
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('輔仁大學智慧校園管理平台')
    run.font.size = Pt(18)
    run.font.color.rgb = FJU_NAVY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('系統分析規格書 (SA Specification)')
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('114 學年度 · 版本 3.0.0 · 2026-04-12')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break
    doc.add_page_break()

    # ══════════════════════════════════════════════
    # VERSION HISTORY
    # ══════════════════════════════════════════════
    doc.add_heading('版本修訂紀錄', level=1)

    version_data = [
        ['版本', '日期', '修訂者', '修訂內容'],
        ['1.0.0', '2026-03-15', '課指組', '初版建立'],
        ['2.0.0', '2026-04-01', '課指組', '新增 AI 模組規格、校園地圖規格'],
        ['2.1.0', '2026-04-08', '課指組', '完善前後端架構、API 清單、測試規範'],
        ['3.0.0', '2026-04-12', '課指組', '更新為 Laravel 12 + SQLite 架構；完整 21 表 Database Schema 詳細欄位定義；更新技術選型與 API 清單；新增修繕追蹤與行事曆管理模組'],
    ]
    tbl = doc.add_table(rows=len(version_data), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(version_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════
    add_toc(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════
    # 1. 系統概述
    # ══════════════════════════════════════════════
    doc.add_heading('1. 系統概述', level=1)

    doc.add_heading('1.1 專案背景', level=2)
    doc.add_paragraph(
        '輔仁大學課外活動指導組（課指組）長期面臨紙本流程繁瑣、場地/器材管理不透明、'
        '社團活動審核耗時等痛點。FJU Smart Hub 旨在透過 AI 與數位化轉型，建構一站式智慧校園管理平台，'
        '涵蓋場地預約、器材借用、活動管理、AI 預審、數位證書等十大核心模組。'
    )

    doc.add_heading('1.2 專案目標', level=2)
    goals = [
        '建構整合前後端的全棧 Web 應用，取代現行紙本流程',
        '導入 AI (Dify RAG) 智慧預審，將審核時間從 3 天縮短至 30 秒',
        '實現三階段資源調度（志願序配對 → AI 協商 → 官方核准）',
        '提供 Leaflet.js 無障礙互動校園地圖（4 LayerGroups）',
        '整合 Google OAuth + 2FA (TOTP/SMS) 全方位安全機制',
        '五國語言支援（繁中、簡中、英文、日文、韓文）',
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')

    doc.add_heading('1.3 適用對象（五大角色）', level=2)
    roles_data = [
        ['角色', '代碼', '權限說明'],
        ['系統管理員', 'admin', '全系統管理、信用積分仲裁、強制介入'],
        ['課指組幹事', 'officer', '審核活動/預約、管理器材、發放證書'],
        ['指導教授', 'professor', '活動指導、預算簽核、成績評鑑'],
        ['學生（社團幹部/一般生）', 'student', '預約場地、申請活動、借用器材、報名活動'],
        ['IT 管理員', 'it', '系統監控、日誌管理、API 維護'],
    ]
    tbl = doc.add_table(rows=len(roles_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(roles_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('1.4 系統範圍', level=2)
    doc.add_paragraph('本系統涵蓋以下十大核心模組（The 10 Pillars）：')
    modules_data = [
        ['#', '模組名稱', '功能概述'],
        ['1', '場地活化大數據', '場地預約、三階段調度、使用率分析'],
        ['2', '器材盤點追蹤', '器材借還、LINE 提醒、信用積分連動'],
        ['3', '行事曆管理', '校園活動行事曆、毛玻璃面板'],
        ['4', '社團資訊', '102 社團/學會資訊、分類篩選'],
        ['5', '動態活動牆', '活動展示、報名、篩選搜尋'],
        ['6', 'AI 智慧預審', 'Dify RAG 法規比對、風險評估'],
        ['7', 'AI 導覽 (輔寶)', '柴犬 AI 助理、法規查詢、流程引導'],
        ['8', 'AI 企劃生成器', 'Dify Workflow 自動生成活動企劃'],
        ['9', '職能 E-Portfolio', '職能雷達圖、活動紀錄、PDF 匯出'],
        ['10', '幹部證書自動化', '數位簽章、QR Code 驗證'],
    ]
    tbl = doc.add_table(rows=len(modules_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(modules_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    # ══════════════════════════════════════════════
    # 2. 技術架構
    # ══════════════════════════════════════════════
    doc.add_heading('2. 技術架構', level=1)

    doc.add_heading('2.1 技術選型', level=2)
    tech_data = [
        ['層級', '技術', '版本', '用途'],
        ['後端框架', 'Laravel', '12.56.0', 'RESTful API、Eloquent ORM、Blade 模板'],
        ['程式語言', 'PHP', '8.2', '後端邏輯、Artisan CLI'],
        ['資料庫', 'SQLite / MySQL 8.0', '-', '結構化資料存儲（開發 SQLite / 正式 MySQL）'],
        ['前端模板', 'Blade + Tailwind CSS', '4.x', '伺服器端渲染、響應式設計'],
        ['地圖引擎', 'Leaflet.js', '1.9.4', '互動校園地圖、GeoJSON 建築分區'],
        ['圖表', 'Chart.js', '4.x', '數據視覺化、雷達圖、長條圖'],
        ['AI 引擎', 'Dify RAG + GPT-4', '-', '法規預審、企劃生成、協商建議'],
        ['安全防護', 'Cloudflare WAF', '-', 'DDoS 防護、IP 過濾、Bot 管理'],
        ['認證', 'Google OAuth + TOTP', '-', '單一登入、雙因子驗證'],
        ['部署', 'PM2 + Artisan Serve', '-', '進程管理、服務守護'],
        ['版控', 'Git + GitHub', '-', '版本控制、協作開發'],
        ['CDN', 'Cloudflare + jsDelivr', '-', '靜態資源加速、前端庫載入'],
        ['推播', 'LINE Notify + SMS', '-', '到期提醒、審核通知'],
        ['儲存', 'Cloudflare R2', '-', '時光膠囊文件加密存放'],
        ['向量 DB', 'Pinecone', '-', 'RAG 法規文件索引'],
        ['國際化', 'i18n JSON', '5 語系', '繁中/簡中/英文/日文/韓文'],
    ]
    tbl = doc.add_table(rows=len(tech_data), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(tech_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('2.2 系統架構圖', level=2)
    doc.add_paragraph(
        '系統採用 MVC 架構，Blade 模板伺服器端渲染，透過 RESTful API 與前端 JavaScript 互動：'
    )
    arch_text = (
        '[使用者] ← HTTPS/Cloudflare WAF →\n'
        '  [Blade 模板 + Tailwind CSS + Leaflet.js]\n'
        '    ↕ REST API (JSON)\n'
        '  [Laravel 12 (PHP 8.2) + Eloquent ORM]\n'
        '    ↕ Database\n'
        '  [SQLite (開發) / MySQL 8.0 (正式)]\n'
        '    ↕ External Services\n'
        '  [Dify RAG | Cloudflare R2 | LINE Notify | Pinecone]'
    )
    p = doc.add_paragraph(arch_text)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    doc.add_heading('2.3 60-30-10 配色規範', level=2)
    color_data = [
        ['比例', '名稱', '色碼', '用途'],
        ['60%', '輔大深藍', '#003153', 'Header、Sidebar、Footer、按鈕'],
        ['30%', '中性灰白', '#F3F4F6 / #FFFFFF', '背景、卡片、內容區'],
        ['10%', '輔大金', '#C49A2A', 'CTA 按鈕、標籤、icon 強調'],
        ['輔助', '功能色', '#28A745 / #DC3545 / #FFC107', '成功/錯誤/警告狀態'],
        ['輔助', '漸層', 'linear-gradient(135deg, #003153, #00507A)', '活動卡片、特色區塊'],
        ['輔助', '毛玻璃', 'rgba(255,255,255,0.85) blur(15px)', '行事曆面板、登入框'],
    ]
    tbl = doc.add_table(rows=len(color_data), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(color_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('2.4 UI/UX 規範', level=2)
    ui_items = [
        '字型：Microsoft JhengHei (微軟正黑體)，備用 system-ui, sans-serif',
        '圓角：12px (fju) / 15px (fju-lg) 統一圓角',
        '佈局：Bento Grid 內部頁面、glassmorphism 登入框',
        'Header：雙層 — 上白 (校名+連結) / 下深藍 (#003153，導覽+搜尋)',
        'Sidebar：深灰 #333333，四大區塊 (主要功能/社群社團/AI核心/管理報表)',
        'Footer：深藍背景，校址、粉專連結、櫃台資訊、版權',
        'AI 聊天：固定右下角 FAB (柴犬 🐕 輔寶)，毛玻璃面板',
        '行事曆面板：毛玻璃效果 (blur 15px, rgba 255,255,255,0.85)，佔 40% 寬度',
    ]
    for item in ui_items:
        doc.add_paragraph(item, style='List Bullet')

    # ══════════════════════════════════════════════
    # 3. 資料庫設計 (COMPLETE 21 TABLES)
    # ══════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('3. 資料庫設計', level=1)

    doc.add_heading('3.1 ER 模型概要', level=2)
    doc.add_paragraph(
        '本系統使用 SQLite（開發環境）/ MySQL 8.0 InnoDB（正式環境）作為主資料庫，'
        '共 21 張資料表，支援外鍵約束、索引優化、事務控制。以下為完整資料表清單及其詳細欄位定義。'
    )

    doc.add_heading('3.2 資料表總覽', level=2)

    summary_data = [
        ['#', '資料表', '中文名稱', '欄位數', '主要用途'],
        ['1', 'users', '使用者', '18', '使用者帳號、角色權限、信用積分、2FA 設定'],
        ['2', 'clubs', '社團', '12', '社團/學會基本資料、分類、指導老師'],
        ['3', 'club_members', '社團成員', '6', '社團成員關聯、職位、加退社時間'],
        ['4', 'venues', '場地', '10', '校園場地資訊、容量、設備、座標'],
        ['5', 'activities', '活動', '19', '社團活動管理、審核、預算、AI 預審'],
        ['6', 'activity_registrations', '活動報名', '6', '活動報名紀錄、簽到、狀態追蹤'],
        ['7', 'equipment', '器材', '10', '器材清冊、編號、狀態、購入資訊'],
        ['8', 'equipment_borrowings', '器材借用', '11', '借用紀錄、歸還追蹤、提醒機制'],
        ['9', 'reservations', '場地預約', '17', '三階段場地預約、志願序、AI 審核'],
        ['10', 'conflicts', '衝突協商', '15', '場地衝突偵測、AI 介入、協商紀錄'],
        ['11', 'credit_logs', '信用紀錄', '8', '信用積分加扣紀錄、原因追蹤'],
        ['12', 'notification_logs', '通知紀錄', '10', '系統/Email/LINE/SMS 通知紀錄'],
        ['13', 'certificates', '證書', '13', '幹部數位證書、簽章、QR 驗證'],
        ['14', 'portfolio_entries', 'E-Portfolio', '9', '職能成長歷程、活動佐證'],
        ['15', 'competency_scores', '職能分數', '4', '五大職能維度分數（雷達圖）'],
        ['16', 'ai_review_logs', 'AI 審核記錄', '11', 'Dify RAG 審核日誌、風險評估'],
        ['17', 'appeals', '申訴', '10', '申訴案件管理、AI 摘要、處理追蹤'],
        ['18', 'repairs', '修繕追蹤', '8', '設備修繕報修、指派、狀態追蹤'],
        ['19', 'calendar_events', '行事曆事件', '8', '校園行事曆事件管理'],
        ['20', 'time_capsules', '時光膠囊', '12', '社團移交文件封裝、R2 儲存'],
        ['21', 'outbox', '異步事件佇列', '6', '事件溯源、異步任務管理'],
    ]
    tbl = doc.add_table(rows=len(summary_data), cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(summary_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(8.5)
                run.font.name = 'Microsoft JhengHei'
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
            elif i % 2 == 0:
                set_cell_shading(cell, ROW_ALT_BG)

    doc.add_heading('3.3 關鍵約束與索引', level=2)
    constraints = [
        'users.role：CHECK IN (admin, officer, professor, student, it)',
        'reservations.priority_level：CHECK IN (1, 2, 3)，Level 1 為校方最高優先',
        'reservations.stage：CHECK IN (algorithm, negotiation, approval, approved, rejected)',
        'credit_logs.action：CHECK IN (add, deduct)，低於 60 分觸發 JWT 失效',
        'venues.status：CHECK IN (available, occupied, maintenance, reserved)',
        '索引：idx_users_role, idx_users_credit_score, idx_activities_club_id, idx_activities_event_date, idx_activities_status, idx_reservations_date, idx_reservations_venue_id, idx_reservations_status, idx_credit_logs_user, idx_notification_logs_user, idx_notification_logs_read, idx_equipment_status, idx_equipment_code, idx_portfolio_user, idx_outbox_status',
        '唯一約束：users.student_id, users.email, equipment.code, certificates.certificate_code, appeals.code, repairs.code, club_members(club_id+user_id), activity_registrations(activity_id+user_id), competency_scores(user_id+dimension)',
    ]
    for c in constraints:
        doc.add_paragraph(c, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('事務控制：')
    run.font.bold = True
    doc.add_paragraph(
        '場地預約採用 Pessimistic Locking (SELECT ... FOR UPDATE) 確保並行預約時無 Race Condition。'
        '信用積分扣抵使用 START TRANSACTION 確保一致性。'
    )

    # ══════════════════════════════════════════════
    # 3.4 完整 21 表欄位定義
    # ══════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('3.4 完整資料表欄位定義（21 Tables）', level=2)
    doc.add_paragraph(
        '以下為系統所有 21 張資料表的完整欄位定義，包含欄位名稱、資料型別、預設值、可為空屬性及說明。'
        '欄位型別中 ENUM 以紅色標示，外鍵參照以藍色標示。'
    )

    # ─── Table 1: users ───
    make_db_detail_table(doc, 1, 'users', '使用者',
        '儲存所有平台使用者資訊，支援五種角色（admin/officer/professor/student/it）、'
        'Google OAuth 登入、TOTP 雙因子驗證、信用積分制度。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵，自動遞增'],
            ['student_id', 'VARCHAR(20) UNIQUE', '-', '否', '學號（唯一）'],
            ['name', 'VARCHAR(100)', '-', '否', '使用者姓名'],
            ['email', 'VARCHAR(191) UNIQUE', '-', '否', 'Email（唯一）'],
            ['phone', 'VARCHAR(30)', '-', '是', '聯絡電話'],
            ['role', 'ENUM(admin, officer, professor, student, it)', "'student'", '否', '角色權限（五種）'],
            ['club_position', 'VARCHAR(100)', '-', '是', '社團職位（社長/副社長等）'],
            ['credit_score', 'INTEGER', '100', '否', '信用積分（起始 100，低於 60 觸發限制）'],
            ['google_oauth_id', 'VARCHAR(255)', '-', '是', 'Google OAuth ID'],
            ['google_avatar_url', 'VARCHAR(255)', '-', '是', 'Google 頭像 URL'],
            ['two_factor_enabled', 'BOOLEAN', 'false', '否', '是否啟用 2FA'],
            ['two_factor_secret', 'VARCHAR(255)', '-', '是', 'TOTP 密鑰'],
            ['avatar_url', 'VARCHAR(255)', '-', '是', '自訂頭像 URL'],
            ['language', 'VARCHAR(10)', "'zh-TW'", '否', '介面語言偏好'],
            ['is_active', 'BOOLEAN', 'true', '否', '帳號是否啟用'],
            ['password', 'VARCHAR(255)', "''", '否', '密碼（OAuth 使用者為空）'],
            ['created_at', 'TIMESTAMP', 'CURRENT', '是', '建立時間'],
            ['updated_at', 'TIMESTAMP', 'CURRENT', '是', '更新時間'],
        ],
        [
            ['索引', 'INDEX(role), INDEX(credit_score)'],
            ['唯一', 'UNIQUE(student_id), UNIQUE(email)'],
        ]
    )

    # ─── Table 2: clubs ───
    make_db_detail_table(doc, 2, 'clubs', '社團',
        '儲存所有社團與學會的基本資料，支援分類篩選、指導老師與社長關聯。系統內建 114 學年度完整資料（社團 67 + 學會 50 = 117 組織，demo 載入 102 筆）。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['name', 'VARCHAR(100)', '-', '否', '社團/學會名稱'],
            ['category', 'VARCHAR(50)', '-', '否', '分類代碼（arts/academic/service/sports/leisure）'],
            ['category_label', 'VARCHAR(50)', '-', '是', '分類中文名稱（學藝/學術/服務/體育/康樂）'],
            ['type', "ENUM('club', 'association')", "'club'", '否', '組織類型（社團/學會）'],
            ['description', 'TEXT', '-', '是', '社團描述'],
            ['advisor_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '指導老師 ID'],
            ['president_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '社長 ID'],
            ['member_count', 'INTEGER', '0', '否', '成員人數'],
            ['established_year', 'VARCHAR(10)', '-', '是', '創社年份'],
            ['is_active', 'BOOLEAN', 'true', '否', '是否活躍'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(category), INDEX(type)'],
        ]
    )

    # ─── Table 3: club_members ───
    make_db_detail_table(doc, 3, 'club_members', '社團成員',
        '社團與使用者的多對多關聯表，記錄成員職位與加退社時間。刪除社團或使用者時級聯刪除。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['club_id', 'BIGINT UNSIGNED (FK → clubs.id)', '-', '否', '所屬社團 ID（CASCADE DELETE）'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '成員 ID（CASCADE DELETE）'],
            ['position', 'VARCHAR(50)', "'member'", '否', '職位（member/president/vp 等）'],
            ['joined_at', 'TIMESTAMP', '-', '是', '加入時間'],
            ['left_at', 'TIMESTAMP', '-', '是', '退社時間'],
        ],
        [
            ['唯一', 'UNIQUE(club_id, user_id)'],
            ['外鍵', 'club_id → clubs.id ON DELETE CASCADE, user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 4: venues ───
    make_db_detail_table(doc, 4, 'venues', '場地',
        '校園場地資訊，含地理座標（Leaflet.js 地圖標記）、設備清單、平面圖。支援四種狀態管理。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['name', 'VARCHAR(100)', '-', '否', '場地名稱（如中美堂、文華樓）'],
            ['location', 'VARCHAR(200)', '-', '是', '位置描述'],
            ['capacity', 'INTEGER', '0', '否', '最大容納人數'],
            ['status', "ENUM('available','occupied','maintenance','reserved')", "'available'", '否', '場地狀態'],
            ['equipment_list', 'TEXT', '-', '是', '場地設備清單（JSON/文字）'],
            ['floor_plan_url', 'VARCHAR(255)', '-', '是', '平面圖 URL'],
            ['latitude', 'DECIMAL(10,6)', '-', '是', '緯度（WGS84）'],
            ['longitude', 'DECIMAL(10,6)', '-', '是', '經度（WGS84）'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(status)'],
        ]
    )

    # ─── Table 5: activities ───
    make_db_detail_table(doc, 5, 'activities', '活動',
        '社團活動完整資訊，涵蓋審核狀態、AI 預審結果、預算申請與核准金額。支援五種狀態流轉。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['title', 'VARCHAR(200)', '-', '否', '活動名稱'],
            ['description', 'TEXT', '-', '是', '活動描述'],
            ['club_id', 'BIGINT UNSIGNED (FK → clubs.id)', '-', '是', '主辦社團 ID'],
            ['venue_id', 'BIGINT UNSIGNED (FK → venues.id)', '-', '是', '活動場地 ID'],
            ['organizer_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '主辦人 ID'],
            ['organizer_name', 'VARCHAR(100)', '-', '是', '主辦人姓名'],
            ['event_date', 'DATE', '-', '否', '活動日期'],
            ['start_time', 'TIME', '-', '否', '開始時間'],
            ['end_time', 'TIME', '-', '否', '結束時間'],
            ['max_participants', 'INTEGER', '0', '否', '最大參與人數'],
            ['current_participants', 'INTEGER', '0', '否', '目前已報名人數'],
            ['category', 'VARCHAR(50)', '-', '是', '活動類別（學術/服務/康樂/體育/藝文）'],
            ['status', "ENUM('pending','approved','rejected','completed','cancelled')", "'pending'", '否', '審核狀態'],
            ['approval_number', 'VARCHAR(255)', '-', '是', '核准文號'],
            ['ai_review_result', 'TEXT', '-', '是', 'AI 預審結果（JSON）'],
            ['budget_requested', 'DECIMAL(12,2)', '0.00', '否', '申請預算'],
            ['budget_approved', 'DECIMAL(12,2)', '0.00', '否', '核准預算'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(club_id), INDEX(event_date), INDEX(status)'],
        ]
    )

    # ─── Table 6: activity_registrations ───
    make_db_detail_table(doc, 6, 'activity_registrations', '活動報名',
        '活動報名紀錄，支援報名、出席、取消三種狀態，以及簽到時間記錄。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['activity_id', 'BIGINT UNSIGNED (FK → activities.id)', '-', '否', '活動 ID（CASCADE DELETE）'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '報名者 ID（CASCADE DELETE）'],
            ['status', "ENUM('registered','attended','cancelled')", "'registered'", '否', '報名狀態'],
            ['check_in_at', 'TIMESTAMP', '-', '是', '簽到時間'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['唯一', 'UNIQUE(activity_id, user_id)'],
            ['外鍵', 'activity_id → activities.id ON DELETE CASCADE, user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 7: equipment ───
    make_db_detail_table(doc, 7, 'equipment', '器材',
        '器材清冊管理，含唯一編號、分類、狀態追蹤、購入成本與圖片。支援四種狀態（可借用/已借出/維修中/已報廢）。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['code', 'VARCHAR(20) UNIQUE', '-', '否', '器材編號（如 EQ-001）'],
            ['name', 'VARCHAR(150)', '-', '否', '器材名稱'],
            ['category', 'VARCHAR(50)', '-', '否', '分類（攝影/音響/照明等）'],
            ['status', "ENUM('available','borrowed','maintenance','retired')", "'available'", '否', '器材狀態'],
            ['condition_note', 'VARCHAR(255)', '-', '是', '狀態備註'],
            ['purchase_date', 'DATE', '-', '是', '購入日期'],
            ['cost', 'DECIMAL(12,2)', '0.00', '否', '購入成本'],
            ['image_url', 'VARCHAR(255)', '-', '是', '器材圖片 URL'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(status), INDEX(code)'],
            ['唯一', 'UNIQUE(code)'],
        ]
    )

    # ─── Table 8: equipment_borrowings ───
    make_db_detail_table(doc, 8, 'equipment_borrowings', '器材借用',
        '器材借用紀錄，追蹤借出/歸還/逾期/遺失狀態，支援 LINE 到期提醒、信用積分連動。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['equipment_id', 'BIGINT UNSIGNED (FK → equipment.id)', '-', '否', '器材 ID（CASCADE DELETE）'],
            ['borrower_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '借用者 ID（CASCADE DELETE）'],
            ['borrower_name', 'VARCHAR(100)', '-', '是', '借用者姓名（冗餘加速查詢）'],
            ['borrow_date', 'DATE', '-', '否', '借出日期'],
            ['expected_return_date', 'DATE', '-', '否', '預期歸還日期'],
            ['actual_return_date', 'DATE', '-', '是', '實際歸還日期'],
            ['status', "ENUM('active','returned','overdue','lost')", "'active'", '否', '借用狀態'],
            ['return_condition', 'TEXT', '-', '是', '歸還時狀態描述'],
            ['reminder_sent', 'BOOLEAN', 'false', '否', '是否已發送到期提醒'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['外鍵', 'equipment_id → equipment.id ON DELETE CASCADE, borrower_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 9: reservations ───
    make_db_detail_table(doc, 9, 'reservations', '場地預約',
        '三階段場地預約系統核心表，支援志願序配對（Algorithm）→ 自主協商（Negotiation）→ 官方核准（Approval）流程。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '申請者 ID'],
            ['venue_id', 'BIGINT UNSIGNED (FK → venues.id)', '-', '否', '場地 ID'],
            ['activity_id', 'BIGINT UNSIGNED (FK → activities.id)', '-', '是', '關聯活動 ID'],
            ['user_name', 'VARCHAR(100)', '-', '是', '申請者姓名（冗餘）'],
            ['club_name', 'VARCHAR(100)', '-', '是', '社團名稱（冗餘）'],
            ['venue_name', 'VARCHAR(100)', '-', '是', '場地名稱（冗餘）'],
            ['reservation_date', 'DATE', '-', '否', '預約日期'],
            ['start_time', 'TIME', '-', '否', '開始時間'],
            ['end_time', 'TIME', '-', '否', '結束時間'],
            ['priority_level', 'INTEGER', '3', '否', '優先等級（1=校方最高/2/3=一般）'],
            ['status', "ENUM('pending','confirmed','negotiating','rejected','cancelled')", "'pending'", '否', '預約狀態'],
            ['stage', "ENUM('algorithm','negotiation','approval','approved','rejected')", "'algorithm'", '否', '流程階段'],
            ['preference_order', 'INTEGER', '1', '否', '志願序（第幾志願）'],
            ['ai_review_result', 'TEXT', '-', '是', 'AI 審核結果（JSON）'],
            ['purpose', 'TEXT', '-', '是', '使用目的'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(reservation_date), INDEX(venue_id), INDEX(status)'],
        ]
    )

    # ─── Table 10: conflicts ───
    make_db_detail_table(doc, 10, 'conflicts', '衝突協商',
        '場地預約衝突偵測與協商管理，支援 AI 自動介入、計時器強制關閉、協商日誌。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['reservation_a_id', 'BIGINT UNSIGNED (FK → reservations.id)', '-', '是', '衝突方 A 預約 ID'],
            ['reservation_b_id', 'BIGINT UNSIGNED (FK → reservations.id)', '-', '是', '衝突方 B 預約 ID'],
            ['venue_id', 'BIGINT UNSIGNED (FK → venues.id)', '-', '是', '場地 ID'],
            ['party_a', 'VARCHAR(100)', '-', '是', '衝突方 A 名稱'],
            ['party_b', 'VARCHAR(100)', '-', '是', '衝突方 B 名稱'],
            ['venue_name', 'VARCHAR(100)', '-', '是', '場地名稱（冗餘）'],
            ['conflict_date', 'DATE', '-', '否', '衝突日期'],
            ['time_slot', 'VARCHAR(50)', '-', '否', '衝突時段'],
            ['status', "ENUM('pending','negotiating','resolved','escalated')", "'pending'", '否', '衝突狀態'],
            ['stage', "ENUM('initial','ai_intervention','forced_close')", "'initial'", '否', '處理階段'],
            ['elapsed_minutes', 'INTEGER', '0', '否', '已經過分鐘數（3 分鐘 AI 介入/6 分鐘強制關閉）'],
            ['negotiation_log', 'TEXT', '-', '是', '協商過程日誌（JSON）'],
            ['ai_suggestion', 'TEXT', '-', '是', 'AI 建議方案（JSON）'],
            ['resolution', 'TEXT', '-', '是', '最終決議'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        []
    )

    # ─── Table 11: credit_logs ───
    make_db_detail_table(doc, 11, 'credit_logs', '信用紀錄',
        '信用積分加扣紀錄，追蹤每次積分變動的原因與關聯對象。積分低於 60 分觸發 JWT 自動失效（Observer Pattern）。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '使用者 ID（CASCADE DELETE）'],
            ['action', "ENUM('add', 'deduct')", '-', '否', '加分/扣分'],
            ['points', 'INTEGER', '-', '否', '積分變動量'],
            ['reason', 'VARCHAR(200)', '-', '否', '變動原因說明'],
            ['reference_type', 'VARCHAR(50)', '-', '是', '關聯類型（reservation/equipment 等）'],
            ['reference_id', 'BIGINT UNSIGNED', '-', '是', '關聯記錄 ID'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(user_id)'],
            ['外鍵', 'user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 12: notification_logs ───
    make_db_detail_table(doc, 12, 'notification_logs', '通知紀錄',
        '多管道通知紀錄，支援系統內通知、Email、LINE Notify、SMS 四種管道，含已讀追蹤。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '接收者 ID（CASCADE DELETE）'],
            ['title', 'VARCHAR(200)', '-', '否', '通知標題'],
            ['message', 'TEXT', '-', '否', '通知內容'],
            ['channel', "ENUM('system','email','line','sms')", "'system'", '否', '通知管道'],
            ['read', 'BOOLEAN', 'false', '否', '是否已讀'],
            ['tracking_token', 'VARCHAR(255)', '-', '是', '追蹤令牌（Email 開啟追蹤）'],
            ['read_at', 'TIMESTAMP', '-', '是', '已讀時間'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(user_id), INDEX(read)'],
            ['外鍵', 'user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 13: certificates ───
    make_db_detail_table(doc, 13, 'certificates', '證書',
        '幹部數位證書自動生成，含 SHA256 數位簽章、唯一驗證碼、QR Code 線上驗證連結、PDF 下載。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '持證人 ID（CASCADE DELETE）'],
            ['club_id', 'BIGINT UNSIGNED (FK → clubs.id)', '-', '是', '社團 ID'],
            ['name', 'VARCHAR(100)', '-', '否', '持證人姓名'],
            ['club_name', 'VARCHAR(100)', '-', '是', '社團名稱（冗餘）'],
            ['position', 'VARCHAR(50)', '-', '否', '幹部職位'],
            ['term', 'VARCHAR(50)', '-', '否', '任期（如 114 學年度上學期）'],
            ['certificate_code', 'VARCHAR(50) UNIQUE', '-', '否', '唯一證書編碼'],
            ['digital_signature', 'VARCHAR(255)', '-', '是', 'SHA256 數位簽章'],
            ['verification_url', 'VARCHAR(255)', '-', '是', 'QR Code 驗證連結'],
            ['issued_at', 'TIMESTAMP', '-', '是', '核發日期'],
            ['pdf_url', 'VARCHAR(255)', '-', '是', 'PDF 下載連結'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['唯一', 'UNIQUE(certificate_code)'],
            ['外鍵', 'user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 14: portfolio_entries ───
    make_db_detail_table(doc, 14, 'portfolio_entries', 'E-Portfolio 紀錄',
        '學生職能成長歷程記錄，支援活動佐證、技能標籤、分類管理，與活動紀錄自動彙總。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '使用者 ID（CASCADE DELETE）'],
            ['title', 'VARCHAR(200)', '-', '否', '紀錄標題'],
            ['description', 'TEXT', '-', '是', '詳細描述'],
            ['category', 'VARCHAR(50)', '-', '是', '分類（活動/競賽/志工等）'],
            ['tags', 'VARCHAR(255)', '-', '是', '技能標籤（逗號分隔）'],
            ['activity_id', 'BIGINT UNSIGNED (FK → activities.id)', '-', '是', '關聯活動 ID'],
            ['evidence_url', 'VARCHAR(255)', '-', '是', '佐證資料 URL'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(user_id)'],
            ['外鍵', 'user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 15: competency_scores ───
    make_db_detail_table(doc, 15, 'competency_scores', '職能分數',
        '五大職能維度分數（領導力、創意思維、團隊合作、溝通能力、數位能力），用於 Chart.js 雷達圖呈現。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['user_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '否', '使用者 ID（CASCADE DELETE）'],
            ['dimension', 'VARCHAR(50)', '-', '否', '維度名稱（leadership/creativity/teamwork/communication/digital）'],
            ['score', 'INTEGER', '0', '否', '分數（0-100）'],
            ['updated_at', 'TIMESTAMP', '-', '是', '最後更新時間'],
        ],
        [
            ['唯一', 'UNIQUE(user_id, dimension)'],
            ['外鍵', 'user_id → users.id ON DELETE CASCADE'],
        ]
    )

    # ─── Table 16: ai_review_logs ───
    make_db_detail_table(doc, 16, 'ai_review_logs', 'AI 審核記錄',
        'Dify RAG 審核日誌，記錄每次 AI 預審的輸入/輸出、風險等級、違規項目、法規參照。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['target_type', 'VARCHAR(50)', '-', '否', '審核對象類型（activity/reservation 等）'],
            ['target_id', 'BIGINT UNSIGNED', '-', '否', '審核對象 ID'],
            ['reviewer_type', 'VARCHAR(50)', "'dify_rag'", '否', '審核引擎（dify_rag/gpt4 等）'],
            ['input_data', 'TEXT', '-', '是', '輸入資料（JSON）'],
            ['output_data', 'TEXT', '-', '是', '輸出結果（JSON）'],
            ['allow_next_step', 'BOOLEAN', 'true', '否', '是否允許進入下一步'],
            ['risk_level', 'VARCHAR(20)', "'Low'", '否', '風險等級（Low/Medium/High）'],
            ['violations', 'TEXT', '-', '是', '違規項目清單（JSON）'],
            ['references', 'TEXT', '-', '是', '參照法規清單（JSON）'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        []
    )

    # ─── Table 17: appeals ───
    make_db_detail_table(doc, 17, 'appeals', '申訴',
        '申訴案件管理，支援場地/器材/活動/其他類型申訴，含 AI 自動摘要、WebSocket 即時推送。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['code', 'VARCHAR(20) UNIQUE', '-', '否', '申訴編號（如 AP-001）'],
            ['appellant_id', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '申訴人 ID'],
            ['appeal_type', 'VARCHAR(50)', '-', '否', '申訴類型（venue/equipment/activity/other）'],
            ['subject', 'VARCHAR(300)', '-', '否', '申訴主旨'],
            ['description', 'TEXT', '-', '是', '申訴內容描述'],
            ['status', "ENUM('pending','processing','resolved','rejected')", "'pending'", '否', '處理狀態'],
            ['ai_summary', 'TEXT', '-', '是', 'AI 摘要（情緒分析/緊急度/建議方案）'],
            ['resolution', 'TEXT', '-', '是', '處理結果'],
            ['assigned_to', 'BIGINT UNSIGNED', '-', '是', '指派處理人 ID'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['唯一', 'UNIQUE(code)'],
        ]
    )

    # ─── Table 18: repairs ───
    make_db_detail_table(doc, 18, 'repairs', '修繕追蹤',
        '設備修繕報修系統，追蹤報修對象、指派維修人員、修繕狀態。支援五種狀態流轉。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['code', 'VARCHAR(20) UNIQUE', '-', '否', '修繕編號（如 RP-001）'],
            ['target', 'VARCHAR(200)', '-', '否', '修繕對象（如投影機、冷氣等）'],
            ['description', 'TEXT', '-', '否', '問題描述'],
            ['status', "ENUM('pending','processing','in_progress','completed','cancelled')", "'pending'", '否', '修繕狀態'],
            ['assignee', 'VARCHAR(100)', '-', '是', '指派維修人員'],
            ['submitted_by', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '報修人 ID'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['唯一', 'UNIQUE(code)'],
        ]
    )

    # ─── Table 19: calendar_events ───
    make_db_detail_table(doc, 19, 'calendar_events', '行事曆事件',
        '校園行事曆事件管理，支援多種事件類型與顏色標記，整合 evo-calendar 前端元件。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['title', 'VARCHAR(200)', '-', '否', '事件標題'],
            ['date', 'DATE', '-', '否', '事件日期'],
            ['type', 'VARCHAR(50)', '-', '是', '事件類型（meeting/deadline/event/holiday 等）'],
            ['color', 'VARCHAR(20)', "'#003153'", '否', '顯示顏色（CSS 色碼）'],
            ['description', 'TEXT', '-', '是', '事件描述'],
            ['venue', 'VARCHAR(100)', '-', '是', '活動場地'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        []
    )

    # ─── Table 20: time_capsules ───
    make_db_detail_table(doc, 20, 'time_capsules', '時光膠囊',
        '社團跨屆移交文件封裝至 Cloudflare R2 加密存放，支援雙重驗證解封、GPS 浮水印防偽。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['club_id', 'BIGINT UNSIGNED (FK → clubs.id)', '-', '是', '社團 ID'],
            ['created_by', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '建立者（前任社長）ID'],
            ['received_by', 'BIGINT UNSIGNED (FK → users.id)', '-', '是', '接收者（現任社長）ID'],
            ['term', 'VARCHAR(50)', '-', '否', '學期（如 114-1）'],
            ['file_manifest', 'TEXT', '-', '是', '文件清單（JSON）'],
            ['r2_storage_key', 'VARCHAR(255)', '-', '是', 'Cloudflare R2 存儲金鑰'],
            ['gps_watermark', 'VARCHAR(255)', '-', '是', 'GPS 浮水印資訊'],
            ['status', "ENUM('sealed','transferred','received')", "'sealed'", '否', '封裝狀態'],
            ['sealed_at', 'TIMESTAMP', '-', '是', '封裝時間'],
            ['transfer_at', 'TIMESTAMP', '-', '是', '移交時間'],
            ['received_at', 'TIMESTAMP', '-', '是', '接收時間'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        []
    )

    # ─── Table 21: outbox ───
    make_db_detail_table(doc, 21, 'outbox', '異步事件佇列',
        '事件溯源（Event Sourcing）設計，管理異步任務佇列，包含 PDF 生成、通知推送、積分計算等背景作業。',
        [
            ['id', 'BIGINT UNSIGNED (PK)', 'AUTO_INCREMENT', '否', '主鍵'],
            ['event_type', 'VARCHAR(100)', '-', '否', '事件類型（pdf_generation/notification 等）'],
            ['payload', 'TEXT', '-', '否', '事件資料（JSON）'],
            ['status', "ENUM('pending','processed','failed')", "'pending'", '否', '處理狀態'],
            ['retry_count', 'INTEGER', '0', '否', '重試次數'],
            ['created_at / updated_at', 'TIMESTAMP', 'CURRENT', '是', '時間戳記'],
        ],
        [
            ['索引', 'INDEX(status)'],
        ]
    )

    # ══════════════════════════════════════════════
    # 3.5 ER 關聯圖（文字描述）
    # ══════════════════════════════════════════════
    doc.add_heading('3.5 ER 關聯摘要', level=2)
    er_relations = [
        'users 1:N → clubs (advisor_id, president_id)',
        'users 1:N → club_members → clubs (M:N 透過 club_members)',
        'users 1:N → activities (organizer_id)',
        'users 1:N → activity_registrations → activities (M:N)',
        'users 1:N → equipment_borrowings → equipment (M:N)',
        'users 1:N → reservations → venues (M:N)',
        'users 1:N → credit_logs',
        'users 1:N → notification_logs',
        'users 1:N → certificates → clubs',
        'users 1:N → portfolio_entries → activities',
        'users 1:N → competency_scores',
        'users 1:N → appeals (appellant_id)',
        'users 1:N → repairs (submitted_by)',
        'users 1:N → time_capsules (created_by, received_by)',
        'clubs 1:N → activities',
        'clubs 1:N → time_capsules',
        'venues 1:N → reservations',
        'venues 1:N → conflicts',
        'reservations 1:N → conflicts (reservation_a_id, reservation_b_id)',
        'activities 1:N → activity_registrations',
    ]
    for r in er_relations:
        doc.add_paragraph(r, style='List Bullet')

    # ══════════════════════════════════════════════
    # 4. 前端架構 (from original doc)
    # ══════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('4. 前端架構', level=1)

    doc.add_heading('4.1 頁面結構', level=2)
    pages_data = [
        ['路由', '頁面名稱', '說明'],
        ['/', '首頁 (Landing)', '公開首頁、功能介紹、登入入口'],
        ['/login', '登入頁', 'Google OAuth + 帳號密碼登入'],
        ['/dashboard', '儀表板', '角色專屬 Dashboard（5 種角色）'],
        ['/campus-map', '校園地圖', 'Leaflet.js 互動地圖 + 行事曆面板'],
        ['/module/venue-booking', '場地預約', '三階段預約 + 使用率分析'],
        ['/module/equipment', '器材管理', '借還管理 + LINE 提醒'],
        ['/module/calendar', '行事曆', '校園行事曆管理'],
        ['/module/club-info', '社團資訊', '102 社團/學會瀏覽'],
        ['/module/activity-wall', '活動牆', '動態活動展示 + 報名'],
        ['/module/ai-overview', 'AI 總覽', 'AI 模組入口'],
        ['/module/ai-guide', 'AI 導覽', '輔寶 AI 助理'],
        ['/module/rag-search', 'RAG 搜尋', '法規智慧問答'],
        ['/module/rag-rental', 'RAG 租借', '場地/器材借用查詢'],
        ['/module/repair', '修繕追蹤', '設備報修管理'],
        ['/module/appeal', '申訴管理', '申訴案件 + AI 摘要'],
        ['/module/reports', '統計報表', '數據分析 Dashboard'],
        ['/module/e-portfolio', '職能 E-Portfolio', '職能雷達圖 + 活動紀錄'],
        ['/module/certificate', '幹部證書', '數位證書生成 + QR 驗證'],
        ['/module/time-capsule', '時光膠囊', '社團移交文件管理'],
        ['/module/2fa', '雙因子驗證', 'TOTP/SMS 設定'],
    ]
    tbl = doc.add_table(rows=len(pages_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(pages_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('4.2 共用元件 (Layout System)', level=2)
    layout_items = [
        'app.blade.php：全局 HTML 骨架，Tailwind CDN、FontAwesome、全局 CSS',
        'shell.blade.php：登入後頁面框架 — 雙層 Header + Sidebar + Main + Footer + Chatbot',
        'Header：上白 (校名/連結/角色切換) + 下深藍 (導覽/搜尋/通知/用戶)',
        'Sidebar：深灰 #333333，四大區塊導覽，active 黃色標記，信用儀表板',
        'Footer：深藍背景，校址、社群連結、聯絡資訊、版權聲明',
        'Chatbot：柴犬 🐕 輔寶 AI 助理 FAB + 對話面板',
    ]
    for item in layout_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.3 互動校園地圖（Leaflet.js）', level=2)
    doc.add_paragraph(
        '地圖以淨心堂 (25.0359, 121.4323) 為中心，使用 OpenStreetMap 底圖 + GeoJSON 建築分區疊加。\n'
        '建築 Popup 包含：學院分色標頭、樓層數、狀態（綠/紅）、「我要預約」按鈕、「定位」按鈕。\n'
        '行事曆面板：毛玻璃效果 (backdrop-filter: blur(15px))，佔 40% 寬度從右側滑入 (GSAP)，整合 evo-calendar 與場地預約狀態。'
    )

    # ══════════════════════════════════════════════
    # 5. 後端架構
    # ══════════════════════════════════════════════
    doc.add_heading('5. 後端架構', level=1)

    doc.add_heading('5.1 API 端點清單', level=2)
    api_data = [
        ['方法', '路徑', '說明'],
        ['GET', '/api/health', '健康檢查（DB/AI/Cache/WAF 狀態）'],
        ['GET', '/api/users', '使用者列表（支援分頁/篩選）'],
        ['POST', '/api/users', '建立使用者'],
        ['GET', '/api/users/{id}', '使用者詳情'],
        ['PUT', '/api/users/{id}', '更新使用者'],
        ['DELETE', '/api/users/{id}', '刪除使用者'],
        ['GET', '/api/clubs', '社團列表（支援 type/category/search 篩選）'],
        ['POST', '/api/clubs', '建立社團'],
        ['GET', '/api/clubs/{id}', '社團詳情'],
        ['PUT', '/api/clubs/{id}', '更新社團'],
        ['DELETE', '/api/clubs/{id}', '刪除社團'],
        ['GET', '/api/clubs/stats', '社團統計（分類/類型分佈）'],
        ['GET', '/api/venues', '場地列表'],
        ['POST', '/api/venues', '建立場地'],
        ['GET', '/api/venues/{id}', '場地詳情'],
        ['PUT', '/api/venues/{id}', '更新場地'],
        ['DELETE', '/api/venues/{id}', '刪除場地'],
        ['GET', '/api/venues/{id}/schedule', '場地時段表（日期查詢）'],
        ['GET', '/api/activities', '活動列表（支援 category/status/search）'],
        ['POST', '/api/activities', '建立活動（含 AI 預審）'],
        ['GET', '/api/activities/{id}', '活動詳情'],
        ['PUT', '/api/activities/{id}', '更新活動'],
        ['DELETE', '/api/activities/{id}', '刪除活動'],
        ['POST', '/api/activities/{id}/register', '報名活動'],
        ['POST', '/api/activities/{id}/cancel-registration', '取消報名'],
        ['GET', '/api/equipment', '器材列表'],
        ['POST', '/api/equipment', '建立器材'],
        ['PUT', '/api/equipment/{id}', '更新器材'],
        ['DELETE', '/api/equipment/{id}', '刪除器材'],
        ['POST', '/api/equipment/borrow', '借用器材'],
        ['POST', '/api/equipment/return', '歸還器材'],
        ['POST', '/api/equipment/remind', '發送到期提醒'],
        ['GET', '/api/reservations', '預約列表（支援 stage/status 篩選）'],
        ['POST', '/api/reservations', '建立預約（含衝突偵測）'],
        ['PUT', '/api/reservations/{id}', '更新預約'],
        ['DELETE', '/api/reservations/{id}', '刪除預約'],
        ['POST', '/api/reservations/{id}/negotiate', 'AI 協商建議'],
        ['POST', '/api/reservations/{id}/accept-suggestion', '接受建議方案'],
        ['GET', '/api/conflicts', '衝突列表'],
        ['POST', '/api/conflicts', '建立衝突記錄'],
        ['POST', '/api/conflicts/{id}/negotiate', '衝突協商'],
        ['GET', '/api/credits/{userId}', '查詢使用者積分'],
        ['POST', '/api/credits/{userId}/deduct', '扣除積分'],
        ['GET', '/api/notifications/{userId}', '通知列表'],
        ['POST', '/api/notifications/{id}/read', '標記已讀'],
        ['POST', '/api/certificates/generate', '生成數位證書'],
        ['GET', '/api/calendar/events', '行事曆事件列表'],
        ['POST', '/api/calendar/events', '建立事件'],
        ['PUT', '/api/calendar/events/{id}', '更新事件'],
        ['DELETE', '/api/calendar/events/{id}', '刪除事件'],
        ['GET', '/api/repairs', '修繕列表'],
        ['POST', '/api/repairs', '提交報修'],
        ['PUT', '/api/repairs/{id}', '更新修繕狀態'],
        ['DELETE', '/api/repairs/{id}', '刪除修繕'],
        ['GET', '/api/appeals', '申訴列表'],
        ['POST', '/api/appeals', '提交申訴'],
        ['PUT', '/api/appeals/{id}', '更新申訴'],
        ['DELETE', '/api/appeals/{id}', '刪除申訴'],
        ['POST', '/api/appeals/{id}/ai-summary', 'AI 申訴摘要'],
        ['GET', '/api/dashboard/stats/{role}', '角色專屬統計'],
        ['POST', '/api/ai/pre-review', 'AI 智慧預審'],
        ['POST', '/api/ai/generate-proposal', 'AI 企劃生成'],
        ['GET', '/api/i18n/{lang}', '多語系資料'],
    ]
    tbl = doc.add_table(rows=len(api_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(api_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(8.5)
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True

    doc.add_heading('5.2 三階段資源調度', level=2)
    doc.add_paragraph('場地預約核心邏輯分三個階段：')
    stages = [
        '第一階段 - 志願序配對（Algorithm）：依 Level 1/2/3 優先權自動配對，無衝突直接進入核准',
        '第二階段 - 自主協商（Negotiation）：偵測衝突後啟動 3/6 分鐘計時器。3 分鐘 GPT-4 介入建議，6 分鐘強制關閉並扣雙方各 10 信用分',
        '第三階段 - 官方核准（Approval）：RAG 法規比對 + Gatekeeping，產出 PDF + TOTP QR 驗證碼',
    ]
    for s in stages:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('5.3 AI 整合（Dify + RAG）', level=2)
    ai_items = [
        'AI 預審：Dify RAG Engine 比對 5+ 法規文件（Pinecone 向量 DB），返回 risk_level / violations / references / suggestions',
        'AI 企劃生成：Dify Workflow 參考法規模板，產出完整企劃書含預算/流程/SDG 對應',
        'AI 協商監控：GPT-4 分析雙方歷史使用紀錄，3 分鐘後自動提供替代方案',
        'AI 申訴摘要：WebSocket 即時推送案件分析，含情緒分析/緊急度/建議方案',
        'AI 導覽 (輔寶)：Shiba-Inu 柴犬化身 FAB，支援法規查詢/場地導覽/流程引導',
    ]
    for item in ai_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.4 安全機制', level=2)
    security_data = [
        ['機制', '技術', '說明'],
        ['WAF', 'Cloudflare WAF', 'DDoS 防護、IP 過濾、Bot 管理'],
        ['認證', 'Google OAuth 2.0', '單一登入、學校 Email 驗證'],
        ['2FA', 'TOTP + SMS', '雙因子驗證（仲裁/核銷/後台必須）'],
        ['加密', 'HTTPS + SHA256', '傳輸加密、證書數位簽章'],
        ['授權', 'Role-based (5 roles)', '角色權限控管'],
        ['信用', 'Credit Score', '積分制度、低分限制、JWT 失效'],
        ['CORS', 'Laravel Middleware', '/api/* 路徑 CORS 開啟'],
        ['輸入驗證', 'Laravel Validation', '伺服器端資料驗證'],
    ]
    tbl = doc.add_table(rows=len(security_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(security_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    # ══════════════════════════════════════════════
    # 6. 核心模組規格 (from original)
    # ══════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('6. 核心模組規格', level=1)

    # 6.1 - 6.10 modules
    modules = [
        {
            'num': '6.1', 'name': '職能 E-Portfolio',
            'desc': '記錄學生參與社團活動的職能成長歷程，提供雷達圖分析與 PDF 履歷匯出。',
            'features': [
                '個人資料表：姓名、學號、Email、社團職位',
                '職能雷達圖（Chart.js radar）：領導力、創意思維、團隊合作、溝通能力、數位能力',
                '活動紀錄列表：自動從 activity_registrations 彙總',
                '技能標籤（Tag-based）：可自訂/自動建議',
                'PDF 匯出：透過 outbox 異步產生',
            ],
            'tables': 'portfolio_entries, competency_scores, activity_registrations',
            'api': 'GET /api/users/:id, GET /api/portfolio/:userId',
        },
        {
            'num': '6.2', 'name': 'AI 企劃生成器',
            'desc': '調用 Dify Workflow 參考校內法規，自動生成完整活動企劃書。',
            'features': [
                '表單輸入：活動名稱、類型、人數、日期、描述',
                'Dify Workflow 調用：GuzzleHTTP → Dify API',
                '生成內容：目的/時間/地點/流程/預算明細/風險評估/SDG 對應',
                'RAG 法規參考列表：5+ 法規文件即時檢索',
                'AI 預審結果：risk_level + confidence 指標',
            ],
            'tables': 'ai_review_logs',
            'api': 'POST /api/ai/generate-proposal, POST /api/ai/pre-review',
        },
        {
            'num': '6.3', 'name': '幹部證書自動化',
            'desc': '自動產生社團幹部數位證書，含數位簽章與 QR Code 驗證。',
            'features': [
                '證書表單：社團名稱、幹部姓名、職位、任期',
                '即時預覽：校內官方格式、金邊裝飾',
                '數位簽章：SHA256 雜湊 + 唯一驗證碼',
                'QR Code：線上驗證連結',
                'PDF 下載：Server-side 產生',
            ],
            'tables': 'certificates',
            'api': 'POST /api/certificates/generate',
        },
        {
            'num': '6.4', 'name': '器材盤點追蹤',
            'desc': '完整器材生命週期管理，含借還流程、LINE/SMS 到期提醒、信用積分連動。',
            'features': [
                '器材清單：編號/名稱/類別/狀態/借用人/歸還日',
                '借用流程：線上申請 → 審核 → 領取 → LINE 提醒 → 歸還',
                '篩選器：全部/可借用/已借出/維修中',
                'LINE Notify：歸還前 1 天自動推送',
                '信用積分：準時歸還 +5，逾期 -2/天',
            ],
            'tables': 'equipment, equipment_borrowings, credit_logs',
            'api': 'GET /api/equipment, POST /api/equipment/borrow, POST /api/equipment/return, POST /api/equipment/remind',
        },
        {
            'num': '6.5', 'name': 'AI 智慧預審 (RAG)',
            'desc': 'Dify RAG Engine 自動比對校內法規，標記合規風險，提供修改建議。',
            'features': [
                '智慧法規問答：自然語言輸入',
                'Pinecone 向量 DB：5+ 法規文件已索引',
                '結構化回覆：引用條文、信心指數、建議列表',
                '常用問題快捷按鈕',
                '租借流程 RAG：場地/器材借用步驟查詢',
            ],
            'tables': 'ai_review_logs',
            'api': 'POST /api/ai/pre-review, GET /api/i18n/:lang',
        },
        {
            'num': '6.6', 'name': '場地活化大數據',
            'desc': '場地使用率熱力圖、周轉排行榜，搭配三階段資源調度系統。',
            'features': [
                '場地統計卡：總數/可用/使用率/維護中',
                '使用熱力圖（Chart.js bar）：時段使用密度',
                '場地周轉排行（Chart.js horizontal bar）',
                '預約表單 Modal：場地/日期/時段/優先等級/用途',
                '三階段調度視覺化：志願序 → 協商 → 核准',
            ],
            'tables': 'venues, reservations, conflicts',
            'api': 'GET /api/venues, POST /api/reservations, POST /api/reservations/:id/negotiate',
        },
        {
            'num': '6.7', 'name': 'AI 申訴摘要',
            'desc': 'WebSocket 即時生成申訴案件摘要，含情緒分析與建議方案。',
            'features': [
                '申訴案件列表：待處理/處理中/已結案',
                '新增申訴表單：主題/描述',
                'AI 摘要面板：案件摘要/建議列表/情緒分析/緊急度',
                '操作按鈕：採納方案/進入仲裁',
                'WebSocket 即時推送（模擬）',
            ],
            'tables': 'appeals',
            'api': 'GET /api/appeals, POST /api/appeals, POST /api/appeals/:id/ai-summary',
        },
        {
            'num': '6.8', 'name': '動態活動牆',
            'desc': '卡片式活動展示，支援分類篩選、關鍵字搜尋、即時報名。',
            'features': [
                '活動卡片：漸層背景/類型標籤/日期/場地/報名進度條',
                '分類篩選：全部/學術/服務/康樂/體育/藝文',
                '關鍵字搜尋：即時過濾',
                '報名功能：POST /api/activities/:id/register',
                '額滿標記：進度條 100% 顯示紅色',
            ],
            'tables': 'activities, activity_registrations',
            'api': 'GET /api/activities, POST /api/activities/:id/register',
        },
        {
            'num': '6.9', 'name': '數位時光膠囊',
            'desc': '社團移交文件封裝至 Cloudflare R2，確保跨屆資產安全傳承。',
            'features': [
                '檔案上傳：R2 Storage 加密存放',
                '文件清單：檔名/大小/上傳日期',
                '封裝狀態：已封裝/已移交/已接收',
                '解封驗證：現任社長 + 前任社長雙重驗證',
                'GPS 浮水印：上傳時記錄地理資訊',
            ],
            'tables': 'time_capsules',
            'api': 'POST /api/capsules, POST /api/capsules/:id/transfer',
        },
        {
            'num': '6.10', 'name': '全方位 2FA',
            'desc': 'TOTP + SMS 雙因子驗證，強制用於仲裁、核銷、後台登入。',
            'features': [
                'TOTP 切換：Google Authenticator / Authy',
                'SMS 切換：手機驗證碼',
                'LINE Notify 綁定：場地/器材通知',
                '必要 2FA 動作：仲裁操作、經費核銷、後台管理登入',
                'QR Code 顯示：TOTP 設定掃描',
            ],
            'tables': 'users (two_factor_enabled, two_factor_secret)',
            'api': 'POST /api/auth/2fa/enable, POST /api/auth/2fa/verify',
        },
    ]

    for m in modules:
        doc.add_heading(f"{m['num']} {m['name']}", level=2)
        doc.add_paragraph(m['desc'])
        doc.add_paragraph('功能需求：')
        for f in m['features']:
            doc.add_paragraph(f, style='List Bullet')
        doc.add_paragraph(f"相關資料表：{m['tables']}")
        doc.add_paragraph(f"API 端點：{m['api']}")

    # ══════════════════════════════════════════════
    # 7. 無障礙校園地圖規格
    # ══════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('7. 無障礙校園地圖規格', level=1)

    doc.add_heading('7.1 技術實現', level=2)
    map_tech = [
        'Leaflet.js v1.9.4 + OpenStreetMap 底圖',
        'GeoJSON 建築分區：campus.geojson (120KB+)，學院分色多邊形',
        '座標系統：WGS84 經緯度，地圖中心 (25.0360, 121.4320) - 淨心堂',
        '淨心堂圓形高亮：L.circle 半徑 50m，金色虛線',
        '兩點仿射變換：中美堂 (A) + 第二圓環 (B) 錨點校正',
    ]
    for item in map_tech:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 四大 LayerGroup', level=2)
    layer_data = [
        ['LayerGroup', '標記數', '標記類型', '顏色', '內容'],
        ['建築標記', '30+', 'Marker', '學院分色', '建築名稱/學院/狀態'],
        ['無障礙設施', '20+', 'Circle', '藍色', '電梯/坡道/無障礙廁所'],
        ['AED 位置', '10+', 'Icon', '紅色', 'AED 設備位置'],
        ['停車場', '5+', 'Polygon', '綠色', '停車區域'],
    ]
    tbl = doc.add_table(rows=len(layer_data), cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(layer_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('7.3 Popup 資訊卡', level=2)
    popup_items = [
        '標頭：學院色彩背景 + 圖示 + 建築名稱 + 英文名',
        '內容：學院歸屬、代碼、樓層數、狀態（綠：可用 / 紅：維護中）',
        '按鈕：「我要預約」(導向 /module/venue-booking) + 「定位」(flyTo zoom 19)',
        'CSS：rounded 12px, box-shadow, 最大寬度 320px',
    ]
    for item in popup_items:
        doc.add_paragraph(item, style='List Bullet')

    # ══════════════════════════════════════════════
    # 8. 導覽結構
    # ══════════════════════════════════════════════
    doc.add_heading('8. 導覽結構', level=1)

    doc.add_heading('8.1 Sidebar 四大區塊', level=2)
    sidebar_data = [
        ['區塊', '項目', '圖示'],
        ['主要功能', '場地預約', 'fa-building'],
        ['', '器材管理', 'fa-tools'],
        ['', '行事曆', 'fa-calendar'],
        ['', '修繕追蹤', 'fa-wrench'],
        ['社群社團', '社團資訊', 'fa-users'],
        ['', '活動牆', 'fa-bullhorn'],
        ['', '時光膠囊', 'fa-box'],
        ['AI 核心', 'AI 總覽', 'fa-brain'],
        ['', 'AI 導覽', 'fa-robot'],
        ['', 'RAG 搜尋', 'fa-search'],
        ['', 'RAG 租借', 'fa-handshake'],
        ['管理報表', '統計報表', 'fa-chart-bar'],
        ['', '申訴管理', 'fa-gavel'],
        ['', 'E-Portfolio', 'fa-id-card'],
        ['', '證書管理', 'fa-certificate'],
    ]
    tbl = doc.add_table(rows=len(sidebar_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(sidebar_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('8.2 信用積分儀表板', level=2)
    doc.add_paragraph(
        'Sidebar 頂部顯示信用積分：起始 100 分、進度條、閾值 60 分警告。'
        '低於 60 分觸發 JWT 自動失效（Observer Pattern）。'
    )

    # ══════════════════════════════════════════════
    # 9. 部署與測試
    # ══════════════════════════════════════════════
    doc.add_heading('9. 部署與測試', level=1)

    doc.add_heading('9.1 Demo 環境', level=2)
    deploy_data = [
        ['項目', '說明'],
        ['Live Demo', 'https://3000-ines7d5od0umg4mb1ae8b-b9b802c4.sandbox.novita.ai'],
        ['GitHub', 'https://github.com/KY0126/SA-v3-'],
        ['框架', 'Laravel 12.56.0 (PHP 8.2)'],
        ['資料庫', 'SQLite (開發) / MySQL 8.0 (正式)'],
        ['進程管理', 'PM2 + artisan serve'],
    ]
    tbl = doc.add_table(rows=len(deploy_data), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(deploy_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    doc.add_heading('9.2 測試涵蓋', level=2)
    tests = [
        '21 頁面路由完整性驗證（curl 測試全部 HTTP 200）',
        '38+ RESTful API 端點回應驗證（GET/POST/PUT/DELETE）',
        'CORS 中間件（/api/* 路徑）',
        '角色切換（admin/officer/professor/student/it）',
        '5 國語系切換（zh-TW/zh-CN/en/ja/ko）',
        '地圖 GeoJSON 載入、4 LayerGroup 切換、建築搜尋/篩選',
        '行事曆面板 GSAP 滑入/滑出動畫',
        '完整 CRUD 測試（users/clubs/venues/activities/equipment/reservations/repairs/appeals/calendar）',
    ]
    for t in tests:
        doc.add_paragraph(t, style='List Bullet')

    # ══════════════════════════════════════════════
    # 10. 附錄
    # ══════════════════════════════════════════════
    doc.add_heading('10. 附錄', level=1)

    doc.add_heading('10.1 社團完整資料 (114 學年度)', level=2)
    doc.add_paragraph(
        '系統內建 114 學年度完整社團資料：社團 67 個 + 學會 50 個 = 共 117 個組織（Demo 載入 102 筆）。'
        '包含名稱、分類、描述、成員數等完整資訊。'
    )

    doc.add_heading('10.2 參考法規文件（RAG 知識庫）', level=2)
    regs = [
        '活動申請辦法 v2.1',
        '場地使用管理規則 v3.0',
        '經費補助要點 v1.5',
        '器材借用辦法 v2.0',
        '安全管理規範 v1.8',
    ]
    for r in regs:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('10.3 專案文件清單', level=2)
    files_data = [
        ['檔案', '路徑', '說明'],
        ['Migration', 'database/migrations/', '21 張資料表完整定義'],
        ['Models', 'app/Models/', '18 個 Eloquent 模型'],
        ['Controllers', 'app/Http/Controllers/Api/', '7 個 API 控制器 + PageController'],
        ['Routes', 'routes/api.php', '38+ RESTful API 端點'],
        ['Views', 'resources/views/', '26 個 Blade 模板（21 頁面 + Layout）'],
        ['Seeder', 'database/seeders/', '102 社團 + 7 使用者 + 完整測試資料'],
        ['Static', 'public/', 'CSS/JS/Images 靜態資源'],
        ['Config', '.env', '環境配置（DB/APP/Mail）'],
        ['PM2', 'ecosystem.config.cjs', '進程管理設定'],
        ['SA 規格書', 'FJU_Smart_Hub_SA_Specification.docx', '本文件'],
        ['Git', '.git/', '版本控制歷史'],
        ['README', 'README.md', '專案說明文件'],
    ]
    tbl = doc.add_table(rows=len(files_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(files_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, HEADER_BG)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)

    # ══════════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════════
    output_path = 'FJU_Smart_Hub_SA_Specification.docx'
    doc.save(output_path)
    print(f'✅ Document saved to {output_path}')
    print(f'   Total paragraphs: {len(doc.paragraphs)}')
    print(f'   Total tables: {len(doc.tables)}')


if __name__ == '__main__':
    build_document()
