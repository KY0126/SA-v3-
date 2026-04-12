#!/usr/bin/env python3
"""
Generate FJU Smart Hub SA Specification Document v3.0
with full Database Schema (21 Tables) and auto-TOC.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

for lvl in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[lvl]
    s.font.name = 'Microsoft JhengHei'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

h1 = doc.styles['Heading 1']
h1.font.size = Pt(18)
h1.font.color.rgb = RGBColor(0, 49, 83)
h1.font.bold = True
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.page_break_before = True

h2 = doc.styles['Heading 2']
h2.font.size = Pt(14)
h2.font.color.rgb = RGBColor(0, 49, 83)
h2.font.bold = True
h2.paragraph_format.space_before = Pt(16)
h2.paragraph_format.space_after = Pt(6)

h3 = doc.styles['Heading 3']
h3.font.size = Pt(12)
h3.font.color.rgb = RGBColor(218, 165, 32)
h3.font.bold = True
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after = Pt(4)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
FJU_BLUE = RGBColor(0, 49, 83)
FJU_GOLD = RGBColor(218, 165, 32)
FJU_BG   = RGBColor(244, 246, 248)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_schema_table(doc, table_name_cn, table_name_en, columns, constraints=None, indexes=None):
    """Add a formatted database table schema."""
    doc.add_heading(f'{table_name_cn} ({table_name_en})', level=3)

    # Column table
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['欄位名稱', '資料型別', '長度/精度', '必填', '預設值', '說明']
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_shading(cell, '003153')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(9)

    for col in columns:
        row = tbl.add_row()
        for i, val in enumerate(col):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                if i == 0:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = FJU_BLUE

    # Set column widths
    widths = [Cm(3.2), Cm(2.8), Cm(2.0), Cm(1.2), Cm(2.2), Cm(5.0)]
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # Constraints & indexes
    if constraints:
        p = doc.add_paragraph()
        r = p.add_run('約束條件：')
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = FJU_BLUE
        for c in constraints:
            bp = doc.add_paragraph(c, style='List Bullet')
            for run in bp.runs:
                run.font.size = Pt(9)

    if indexes:
        p = doc.add_paragraph()
        r = p.add_run('索引：')
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = FJU_BLUE
        for idx in indexes:
            bp = doc.add_paragraph(idx, style='List Bullet')
            for run in bp.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()  # spacing


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('天主教輔仁大學')
r.font.size = Pt(16)
r.font.color.rgb = FJU_BLUE
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('學務處 課外活動指導組')
r.font.size = Pt(12)
r.font.color.rgb = FJU_GOLD

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FJU Smart Hub')
r.font.size = Pt(36)
r.font.color.rgb = FJU_BLUE
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('輔仁大學智慧校園管理平台')
r.font.size = Pt(18)
r.font.color.rgb = FJU_GOLD

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('系統分析規格書 (SA Specification)')
r.font.size = Pt(14)
r.font.color.rgb = FJU_BLUE
r.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('114 學年度 · 版本 3.0.1 · 2026-04-12')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(128, 128, 128)

# ============================================================
# VERSION HISTORY (no page break)
# ============================================================
p = doc.add_paragraph()
p.style = doc.styles['Heading 1']
p.text = '版本修訂紀錄'
p.paragraph_format.page_break_before = True

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
for i, h in enumerate(['版本', '日期', '修訂內容', '作者']):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(10)

versions = [
    ['1.0.0', '2026-03-15', '初版 - Hono + Cloudflare Workers 架構', 'SA Team'],
    ['2.0.0', '2026-04-01', '新增 D1 資料庫、10 頁面、28 API', 'SA Team'],
    ['2.1.0', '2026-04-08', '校園地圖 GeoJSON、114 社團資料、行事曆', 'SA Team'],
    ['3.0.0', '2026-04-11', 'Laravel 12 + MySQL 全面改寫，21 表、38+ API、Full CRUD', 'SA Team'],
    ['3.0.1', '2026-04-12', '修復 repairs enum、更新完整 DB Schema 文件', 'SA Team'],
]
for v in versions:
    row = tbl.add_row()
    for i, val in enumerate(v):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(10)

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('目錄', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('（請在 Word 中選取此區域，按 Ctrl+A 全選後按 F9，或右鍵 > 更新功能變數 以產生自動目錄）')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(128, 128, 128)
r.font.italic = True

# Insert TOC field
fld_xml = (
    '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '  <w:sdtPr><w:docPartObj><w:docPartGallery w:val="Table of Contents"/>'
    '  <w:docPartUnique/></w:docPartObj></w:sdtPr>'
    '  <w:sdtContent>'
    '    <w:p><w:r><w:fldChar w:fldCharType="begin"/></w:r>'
    '    <w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
    '    <w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    '    <w:r><w:t>（按 F9 或右鍵更新功能變數以顯示目錄）</w:t></w:r>'
    '    <w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    '  </w:sdtContent>'
    '</w:sdt>'
)
doc.element.body.append(parse_xml(fld_xml))

# ============================================================
# 1. SYSTEM OVERVIEW
# ============================================================
doc.add_heading('1. 系統概述', level=1)

doc.add_heading('1.1 專案背景', level=2)
doc.add_paragraph(
    '輔仁大學課外活動指導組（課指組）長期面臨紙本流程繁瑣、場地/器材管理不透明、社團活動審核耗時等痛點。'
    'FJU Smart Hub 旨在透過 AI 與數位化轉型，建構一站式智慧校園管理平台，涵蓋場地預約、器材借用、活動管理、'
    'AI 預審、數位證書、互動校園地圖等功能，服務全校師生五大角色。'
)

doc.add_heading('1.2 專案目標', level=2)
goals = [
    '建構整合前後端的全棧 Web 應用，取代現行紙本流程',
    '導入 AI (Dify RAG) 智慧預審，將審核時間從 3 天縮短至 30 秒',
    '實現三階段資源調度（志願序配對 → AI 協商 → 官方核准）',
    '提供 Leaflet.js 無障礙互動校園地圖（4 LayerGroups）',
    '整合 Google OAuth + 2FA (TOTP/SMS) 全方位安全機制',
    '五國語言支援（繁中、簡中、英文、日文、韓文）',
    '完整 CRUD 操作支援所有資料實體（21 張資料表、38+ API 端點）',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('1.3 適用對象（五大角色）', level=2)
roles_tbl = doc.add_table(rows=1, cols=4)
roles_tbl.style = 'Table Grid'
for i, h in enumerate(['角色代碼', '角色名稱', '圖示', '權限說明']):
    cell = roles_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(10)

roles = [
    ['admin', '課指組/處室', 'fa-user-tie', '全系統管理、活動審核、報表查看'],
    ['officer', '社團幹部', 'fa-user-shield', '社團管理、活動建立、預算管理'],
    ['professor', '指導教授', 'fa-chalkboard-teacher', '社團指導、風險預警、成效評估'],
    ['student', '一般學生', 'fa-user-graduate', '活動參與、E-Portfolio、證書申請'],
    ['it', '資訊中心', 'fa-server', '系統監控、API 指標、WAF 管理'],
]
for rv in roles:
    row = roles_tbl.add_row()
    for i, val in enumerate(rv):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(10)

doc.add_heading('1.4 系統範圍', level=2)
doc.add_paragraph('本系統涵蓋以下十大核心模組（The 10 Pillars）：')
pillars = [
    '職能 E-Portfolio', 'AI 企劃生成器', '幹部證書自動化', '器材盤點追蹤',
    'AI 智慧預審 (RAG)', '場地活化大數據', 'AI 申訴摘要', '動態活動牆',
    '數位時光膠囊', '全方位 2FA'
]
for i, p_name in enumerate(pillars):
    doc.add_paragraph(f'{i+1}. {p_name}', style='List Bullet')

# ============================================================
# 2. TECH ARCHITECTURE
# ============================================================
doc.add_heading('2. 技術架構', level=1)

doc.add_heading('2.1 技術選型', level=2)
tech_tbl = doc.add_table(rows=1, cols=3)
tech_tbl.style = 'Table Grid'
for i, h in enumerate(['層級', '技術', '版本']):
    cell = tech_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(10)

techs = [
    ['後端框架', 'PHP + Laravel', '8.2 / 12.56.0'],
    ['資料庫', 'MySQL (正式) / SQLite (開發)', '8.0 / 3.x'],
    ['前端 CSS', 'Tailwind CSS (CDN)', '3.x'],
    ['前端模板', 'Laravel Blade Templates', '-'],
    ['圖表', 'Chart.js', '4.4.0'],
    ['地圖', 'Leaflet.js + OpenStreetMap', '1.9.4'],
    ['動畫', 'GSAP + ScrollTrigger', '3.12.5'],
    ['行事曆', 'evo-calendar (jQuery)', '1.1.3'],
    ['圖示', 'FontAwesome Free', '6.4.0'],
    ['AI 引擎', 'Dify RAG (Pinecone)', '-'],
    ['安全', 'Cloudflare WAF + TOTP 2FA', '-'],
    ['通知', 'LINE Notify + SMS + Email', '-'],
]
for t in techs:
    row = tech_tbl.add_row()
    for i, val in enumerate(t):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(10)

doc.add_heading('2.2 系統架構圖', level=2)
doc.add_paragraph(
    '[使用者] ← HTTPS/Cloudflare WAF →\n'
    '  [Tailwind CSS + Blade Templates + Leaflet.js + Chart.js]\n'
    '    ↕ REST API (JSON)\n'
    '  [Laravel 12 (PHP 8.2) — Eloquent ORM]\n'
    '    ↕ SQL\n'
    '  [MySQL 8.0 / SQLite (dev)] — 21 Tables\n'
    '    ↕ HTTP\n'
    '  [Dify RAG (AI)] / [LINE Notify] / [R2 Storage]'
)

doc.add_heading('2.3 60-30-10 配色規範', level=2)
color_tbl = doc.add_table(rows=1, cols=4)
color_tbl.style = 'Table Grid'
for i, h in enumerate(['比例', '用途', '色碼', '名稱']):
    cell = color_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(10)

colors = [
    ['60%', '主色（背景）', '#FFFFFF / #F4F6F8', '白色 / 淺灰'],
    ['30%', '輔色（導覽/標題）', '#003153 / #333333', 'FJU 深藍 / 側邊欄深灰'],
    ['10%', '強調色（按鈕/CTA）', '#DAA520 / #FDB913', 'FJU 金色 / 亮金'],
    ['-', '成功', '#008000', '綠色'],
    ['-', '危險/警告', '#FF0000', '紅色'],
]
for c in colors:
    row = color_tbl.add_row()
    for i, val in enumerate(c):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(10)

doc.add_heading('2.4 UI/UX 規範', level=2)
ui_rules = [
    '字型：Microsoft JhengHei (微軟正黑體)，備用 system-ui, sans-serif',
    '圓角：12px (fju) / 15px (fju-lg) 統一圓角',
    '佈局：Bento Grid 內部頁面、glassmorphism 登入框',
    'Header：雙層 — 上白 (校名+連結) / 下深藍 (#003153，導覽+搜尋)',
    'Sidebar：深灰 #333333，四大區塊 (主要功能/社群社團/AI核心/管理報表)',
    'Footer：深藍背景，校址、粉專連結、櫃台資訊、版權聲明',
    'AI 聊天：固定右下角 FAB (柴犬 輔寶)，毛玻璃面板',
    '行事曆面板：毛玻璃效果 (blur 15px, rgba 255,255,255,0.85)',
]
for u in ui_rules:
    doc.add_paragraph(u, style='List Bullet')

# ============================================================
# 3. DATABASE DESIGN (FULL 21 TABLES)
# ============================================================
doc.add_heading('3. 資料庫設計', level=1)

doc.add_heading('3.1 ER 模型概要', level=2)
doc.add_paragraph(
    '本系統使用 MySQL 8.0 (InnoDB) 作為正式資料庫，開發環境使用 SQLite。'
    '共 21 張資料表，支援外鍵約束、索引優化、事務控制。'
    '以下為完整的資料庫綱要定義。'
)

doc.add_heading('3.2 資料表總覽', level=2)
overview_tbl = doc.add_table(rows=1, cols=4)
overview_tbl.style = 'Table Grid'
for i, h in enumerate(['編號', '資料表名稱', '中文名稱', '欄位數']):
    cell = overview_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(10)

overview_data = [
    ['1','users','使用者','17'],['2','clubs','社團','12'],['3','club_members','社團成員','6'],
    ['4','venues','場地','10'],['5','activities','活動','18'],['6','activity_registrations','活動報名','6'],
    ['7','equipment','器材','10'],['8','equipment_borrowings','器材借用','10'],['9','reservations','場地預約','16'],
    ['10','conflicts','衝突記錄','14'],['11','credit_logs','信用積分記錄','8'],['12','notification_logs','通知記錄','9'],
    ['13','certificates','數位證書','12'],['14','portfolio_entries','E-Portfolio','9'],
    ['15','competency_scores','職能分數','4'],['16','ai_review_logs','AI 審查記錄','10'],
    ['17','appeals','申訴記錄','10'],['18','repairs','報修記錄','8'],
    ['19','calendar_events','行事曆事件','8'],['20','time_capsules','時光膠囊','11'],
    ['21','outbox','事件發送佇列','6'],
]
for d in overview_data:
    row = overview_tbl.add_row()
    for i, val in enumerate(d):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(10)
            if i == 0:
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.3 完整資料表定義', level=2)

# ---- TABLE 1: users ----
add_schema_table(doc, '表 1：使用者', 'users', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵 (AUTO_INCREMENT)'],
    ['student_id', 'VARCHAR', '20', 'Y', '-', '學號/員工編號 (UNIQUE)'],
    ['name', 'VARCHAR', '100', 'Y', '-', '姓名'],
    ['email', 'VARCHAR', '191', 'Y', '-', '電子郵件 (UNIQUE)'],
    ['phone', 'VARCHAR', '30', 'N', 'NULL', '電話號碼'],
    ['role', 'ENUM', '-', 'Y', 'student', "角色：admin, officer, professor, student, it"],
    ['club_position', 'VARCHAR', '100', 'N', 'NULL', '社團職位描述'],
    ['credit_score', 'INT', '-', 'Y', '100', '信用積分 (起始100，低於60強制登出)'],
    ['google_oauth_id', 'VARCHAR', '255', 'N', 'NULL', 'Google OAuth 識別碼'],
    ['google_avatar_url', 'VARCHAR', '255', 'N', 'NULL', 'Google 頭像 URL'],
    ['two_factor_enabled', 'BOOLEAN', '-', 'Y', 'false', '是否啟用 2FA'],
    ['two_factor_secret', 'VARCHAR', '255', 'N', 'NULL', 'TOTP 密鑰'],
    ['avatar_url', 'VARCHAR', '255', 'N', 'NULL', '自訂頭像 URL'],
    ['language', 'VARCHAR', '10', 'Y', 'zh-TW', '語系偏好'],
    ['is_active', 'BOOLEAN', '-', 'Y', 'true', '帳號啟用狀態'],
    ['password', 'VARCHAR', '255', 'Y', "''", '密碼雜湊'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (student_id)',
    'UNIQUE (email)',
    "CHECK role IN ('admin','officer','professor','student','it')",
    'credit_score < 60 → 觸發 JWT 失效 (Observer Pattern)',
], indexes=[
    'INDEX idx_users_role (role)',
    'INDEX idx_users_credit_score (credit_score)',
])

# ---- TABLE 2: clubs ----
add_schema_table(doc, '表 2：社團', 'clubs', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['name', 'VARCHAR', '100', 'Y', '-', '社團名稱'],
    ['category', 'VARCHAR', '50', 'Y', '-', '分類代碼 (academic/recreation/sports/arts/service/religious/social)'],
    ['category_label', 'VARCHAR', '50', 'N', 'NULL', '分類中文標籤 (學藝/康樂/體育/藝術/服務/宗教/聯誼/學會)'],
    ['type', 'ENUM', '-', 'Y', 'club', "類型：club (社團) / association (學會)"],
    ['description', 'TEXT', '-', 'N', 'NULL', '社團描述'],
    ['advisor_id', 'BIGINT U', '-', 'N', 'NULL', '指導老師 ID (FK→users)'],
    ['president_id', 'BIGINT U', '-', 'N', 'NULL', '社長 ID (FK→users)'],
    ['member_count', 'INT', '-', 'Y', '0', '成員數量'],
    ['established_year', 'VARCHAR', '10', 'N', 'NULL', '成立學年度'],
    ['is_active', 'BOOLEAN', '-', 'Y', 'true', '是否活躍'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK type IN ('club','association')",
], indexes=[
    'INDEX idx_clubs_category (category)',
    'INDEX idx_clubs_type (type)',
])

# ---- TABLE 3: club_members ----
add_schema_table(doc, '表 3：社團成員', 'club_members', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['club_id', 'BIGINT U', '-', 'Y', '-', '社團 ID (FK→clubs, CASCADE)'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '使用者 ID (FK→users, CASCADE)'],
    ['position', 'VARCHAR', '50', 'Y', 'member', '職位 (社長/副社長/幹部/社員)'],
    ['joined_at', 'TIMESTAMP', '-', 'N', 'NULL', '加入時間'],
    ['left_at', 'TIMESTAMP', '-', 'N', 'NULL', '離開時間'],
], constraints=[
    'UNIQUE (club_id, user_id)',
    'FOREIGN KEY (club_id) → clubs(id) ON DELETE CASCADE',
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
])

# ---- TABLE 4: venues ----
add_schema_table(doc, '表 4：場地', 'venues', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['name', 'VARCHAR', '100', 'Y', '-', '場地名稱 (中美堂/SF 134/體育館...)'],
    ['location', 'VARCHAR', '200', 'N', 'NULL', '位置描述'],
    ['capacity', 'INT', '-', 'Y', '0', '容納人數'],
    ['status', 'ENUM', '-', 'Y', 'available', "狀態：available/occupied/maintenance/reserved"],
    ['equipment_list', 'TEXT', '-', 'N', 'NULL', '場地設備清單'],
    ['floor_plan_url', 'VARCHAR', '255', 'N', 'NULL', '平面圖 URL'],
    ['latitude', 'DECIMAL', '10,6', 'N', 'NULL', '緯度 (WGS84)'],
    ['longitude', 'DECIMAL', '10,6', 'N', 'NULL', '經度 (WGS84)'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK status IN ('available','occupied','maintenance','reserved')",
], indexes=[
    'INDEX idx_venues_status (status)',
])

# ---- TABLE 5: activities ----
add_schema_table(doc, '表 5：活動', 'activities', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['title', 'VARCHAR', '200', 'Y', '-', '活動名稱'],
    ['description', 'TEXT', '-', 'N', 'NULL', '活動描述'],
    ['club_id', 'BIGINT U', '-', 'N', 'NULL', '主辦社團 ID (FK→clubs)'],
    ['venue_id', 'BIGINT U', '-', 'N', 'NULL', '場地 ID (FK→venues)'],
    ['organizer_id', 'BIGINT U', '-', 'N', 'NULL', '負責人 ID (FK→users)'],
    ['organizer_name', 'VARCHAR', '100', 'N', 'NULL', '負責人姓名'],
    ['event_date', 'DATE', '-', 'Y', '-', '活動日期'],
    ['start_time', 'TIME', '-', 'Y', '-', '開始時間'],
    ['end_time', 'TIME', '-', 'Y', '-', '結束時間'],
    ['max_participants', 'INT', '-', 'Y', '0', '最大參加人數'],
    ['current_participants', 'INT', '-', 'Y', '0', '目前報名人數'],
    ['category', 'VARCHAR', '50', 'N', 'NULL', '活動分類 (academic/service/recreation/sports/arts)'],
    ['status', 'ENUM', '-', 'Y', 'pending', "狀態：pending/approved/rejected/completed/cancelled"],
    ['approval_number', 'VARCHAR', '255', 'N', 'NULL', '核准文號'],
    ['ai_review_result', 'TEXT', '-', 'N', 'NULL', 'AI 預審結果 (JSON)'],
    ['budget_requested', 'DECIMAL', '12,2', 'Y', '0.00', '申請預算'],
    ['budget_approved', 'DECIMAL', '12,2', 'Y', '0.00', '核准預算'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK status IN ('pending','approved','rejected','completed','cancelled')",
], indexes=[
    'INDEX idx_activities_club_id (club_id)',
    'INDEX idx_activities_event_date (event_date)',
    'INDEX idx_activities_status (status)',
])

# ---- TABLE 6: activity_registrations ----
add_schema_table(doc, '表 6：活動報名', 'activity_registrations', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['activity_id', 'BIGINT U', '-', 'Y', '-', '活動 ID (FK→activities, CASCADE)'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '使用者 ID (FK→users, CASCADE)'],
    ['status', 'ENUM', '-', 'Y', 'registered', "狀態：registered/attended/cancelled"],
    ['check_in_at', 'TIMESTAMP', '-', 'N', 'NULL', '簽到時間'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (activity_id, user_id)',
    'FOREIGN KEY (activity_id) → activities(id) ON DELETE CASCADE',
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
])

# ---- TABLE 7: equipment ----
add_schema_table(doc, '表 7：器材', 'equipment', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['code', 'VARCHAR', '20', 'Y', '-', '器材編號 (UNIQUE, e.g. EQ-001)'],
    ['name', 'VARCHAR', '150', 'Y', '-', '器材名稱'],
    ['category', 'VARCHAR', '50', 'Y', '-', '分類 (攝影/音響/視聽/電腦/文具)'],
    ['status', 'ENUM', '-', 'Y', 'available', "狀態：available/borrowed/maintenance/retired"],
    ['condition_note', 'VARCHAR', '255', 'N', 'NULL', '狀態備註'],
    ['purchase_date', 'DATE', '-', 'N', 'NULL', '購入日期'],
    ['cost', 'DECIMAL', '12,2', 'Y', '0.00', '購入成本'],
    ['image_url', 'VARCHAR', '255', 'N', 'NULL', '器材圖片 URL'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (code)',
    "CHECK status IN ('available','borrowed','maintenance','retired')",
], indexes=[
    'INDEX idx_equipment_status (status)',
    'INDEX idx_equipment_code (code)',
])

# ---- TABLE 8: equipment_borrowings ----
add_schema_table(doc, '表 8：器材借用', 'equipment_borrowings', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['equipment_id', 'BIGINT U', '-', 'Y', '-', '器材 ID (FK→equipment, CASCADE)'],
    ['borrower_id', 'BIGINT U', '-', 'Y', '-', '借用者 ID (FK→users, CASCADE)'],
    ['borrower_name', 'VARCHAR', '100', 'N', 'NULL', '借用者姓名 (冗餘快取)'],
    ['borrow_date', 'DATE', '-', 'Y', '-', '借出日期'],
    ['expected_return_date', 'DATE', '-', 'Y', '-', '預計歸還日期'],
    ['actual_return_date', 'DATE', '-', 'N', 'NULL', '實際歸還日期'],
    ['status', 'ENUM', '-', 'Y', 'active', "狀態：active/returned/overdue/lost"],
    ['return_condition', 'TEXT', '-', 'N', 'NULL', '歸還狀況描述'],
    ['reminder_sent', 'BOOLEAN', '-', 'Y', 'false', '是否已發送提醒'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'FOREIGN KEY (equipment_id) → equipment(id) ON DELETE CASCADE',
    'FOREIGN KEY (borrower_id) → users(id) ON DELETE CASCADE',
    "CHECK status IN ('active','returned','overdue','lost')",
])

# ---- TABLE 9: reservations ----
add_schema_table(doc, '表 9：場地預約', 'reservations', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '申請者 ID (FK→users)'],
    ['venue_id', 'BIGINT U', '-', 'Y', '-', '場地 ID (FK→venues)'],
    ['activity_id', 'BIGINT U', '-', 'N', 'NULL', '關聯活動 ID (FK→activities)'],
    ['user_name', 'VARCHAR', '100', 'N', 'NULL', '申請者姓名 (快取)'],
    ['club_name', 'VARCHAR', '100', 'N', 'NULL', '社團名稱 (快取)'],
    ['venue_name', 'VARCHAR', '100', 'N', 'NULL', '場地名稱 (快取)'],
    ['reservation_date', 'DATE', '-', 'Y', '-', '預約日期'],
    ['start_time', 'TIME', '-', 'Y', '-', '開始時間'],
    ['end_time', 'TIME', '-', 'Y', '-', '結束時間'],
    ['priority_level', 'INT', '-', 'Y', '3', '優先等級 (L1校方/L2幹部/L3一般)'],
    ['status', 'ENUM', '-', 'Y', 'pending', "狀態：pending/confirmed/negotiating/rejected/cancelled"],
    ['stage', 'ENUM', '-', 'Y', 'algorithm', "階段：algorithm/negotiation/approval/approved/rejected"],
    ['preference_order', 'INT', '-', 'Y', '1', '志願序'],
    ['ai_review_result', 'TEXT', '-', 'N', 'NULL', 'AI 審查結果 (JSON)'],
    ['purpose', 'TEXT', '-', 'N', 'NULL', '借用目的'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK status IN ('pending','confirmed','negotiating','rejected','cancelled')",
    "CHECK stage IN ('algorithm','negotiation','approval','approved','rejected')",
    'CHECK priority_level IN (1, 2, 3)，Level 1 校方最高優先',
    '場地預約採 Pessimistic Locking (SELECT...FOR UPDATE)',
], indexes=[
    'INDEX idx_reservations_date (reservation_date)',
    'INDEX idx_reservations_venue (venue_id)',
    'INDEX idx_reservations_status (status)',
])

# ---- TABLE 10: conflicts ----
add_schema_table(doc, '表 10：衝突記錄', 'conflicts', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['reservation_a_id', 'BIGINT U', '-', 'N', 'NULL', '預約 A (FK→reservations)'],
    ['reservation_b_id', 'BIGINT U', '-', 'N', 'NULL', '預約 B (FK→reservations)'],
    ['venue_id', 'BIGINT U', '-', 'N', 'NULL', '衝突場地 ID (FK→venues)'],
    ['party_a', 'VARCHAR', '100', 'N', 'NULL', '甲方 (社團名稱)'],
    ['party_b', 'VARCHAR', '100', 'N', 'NULL', '乙方 (社團名稱)'],
    ['venue_name', 'VARCHAR', '100', 'N', 'NULL', '場地名稱 (快取)'],
    ['conflict_date', 'DATE', '-', 'Y', '-', '衝突日期'],
    ['time_slot', 'VARCHAR', '50', 'Y', '-', '衝突時段 (e.g. 14:00-17:00)'],
    ['status', 'ENUM', '-', 'Y', 'pending', "狀態：pending/negotiating/resolved/escalated"],
    ['stage', 'ENUM', '-', 'Y', 'initial', "階段：initial/ai_intervention/forced_close"],
    ['elapsed_minutes', 'INT', '-', 'Y', '0', '已耗時（分鐘），3分鐘AI介入，6分鐘強制關閉'],
    ['negotiation_log', 'TEXT', '-', 'N', 'NULL', '協商紀錄 (JSON)'],
    ['ai_suggestion', 'TEXT', '-', 'N', 'NULL', 'AI 建議方案 (JSON)'],
    ['resolution', 'TEXT', '-', 'N', 'NULL', '最終解決方案'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK status IN ('pending','negotiating','resolved','escalated')",
    "CHECK stage IN ('initial','ai_intervention','forced_close')",
    '3分鐘 → GPT-4 介入建議；6分鐘 → 強制關閉並扣雙方各 10 信用分',
])

# ---- TABLE 11: credit_logs ----
add_schema_table(doc, '表 11：信用積分記錄', 'credit_logs', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '使用者 ID (FK→users, CASCADE)'],
    ['action', 'ENUM', '-', 'Y', '-', "操作：add (加分) / deduct (扣分)"],
    ['points', 'INT', '-', 'Y', '-', '異動點數 (正=加分，負=扣分)'],
    ['reason', 'VARCHAR', '200', 'Y', '-', '異動原因'],
    ['reference_type', 'VARCHAR', '50', 'N', 'NULL', '關聯類型 (reservation/equipment/activity)'],
    ['reference_id', 'BIGINT U', '-', 'N', 'NULL', '關聯記錄 ID'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK action IN ('add','deduct')",
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
], indexes=[
    'INDEX idx_credit_logs_user (user_id)',
])

# ---- TABLE 12: notification_logs ----
add_schema_table(doc, '表 12：通知記錄', 'notification_logs', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '收件者 ID (FK→users, CASCADE)'],
    ['title', 'VARCHAR', '200', 'Y', '-', '通知標題'],
    ['message', 'TEXT', '-', 'Y', '-', '通知內容'],
    ['channel', 'ENUM', '-', 'Y', 'system', "頻道：system/email/line/sms"],
    ['read', 'BOOLEAN', '-', 'Y', 'false', '是否已讀'],
    ['tracking_token', 'VARCHAR', '255', 'N', 'NULL', '追蹤令牌 (用於 Email 開啟追蹤)'],
    ['read_at', 'TIMESTAMP', '-', 'N', 'NULL', '已讀時間'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK channel IN ('system','email','line','sms')",
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
], indexes=[
    'INDEX idx_notification_logs_user (user_id)',
    'INDEX idx_notification_logs_read (read)',
])

# ---- TABLE 13: certificates ----
add_schema_table(doc, '表 13：數位證書', 'certificates', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '持有者 ID (FK→users, CASCADE)'],
    ['club_id', 'BIGINT U', '-', 'N', 'NULL', '社團 ID (FK→clubs)'],
    ['name', 'VARCHAR', '100', 'Y', '-', '幹部姓名'],
    ['club_name', 'VARCHAR', '100', 'N', 'NULL', '社團名稱 (快取)'],
    ['position', 'VARCHAR', '50', 'Y', '-', '職位 (社長/副社長/幹部)'],
    ['term', 'VARCHAR', '50', 'Y', '-', '任期 (e.g. 114學年度第一學期)'],
    ['certificate_code', 'VARCHAR', '50', 'Y', '-', '證書驗證碼 (UNIQUE, e.g. FJU-CERT-2026-000001)'],
    ['digital_signature', 'VARCHAR', '255', 'N', 'NULL', 'SHA256 數位簽章'],
    ['verification_url', 'VARCHAR', '255', 'N', 'NULL', '線上驗證網址'],
    ['issued_at', 'TIMESTAMP', '-', 'N', 'NULL', '核發時間'],
    ['pdf_url', 'VARCHAR', '255', 'N', 'NULL', 'PDF 檔案 URL'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (certificate_code)',
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
])

# ---- TABLE 14: portfolio_entries ----
add_schema_table(doc, '表 14：E-Portfolio 作品集', 'portfolio_entries', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '擁有者 ID (FK→users, CASCADE)'],
    ['title', 'VARCHAR', '200', 'Y', '-', '作品標題'],
    ['description', 'TEXT', '-', 'N', 'NULL', '描述'],
    ['category', 'VARCHAR', '50', 'N', 'NULL', '分類 (活動參與/幹部經歷/競賽成果)'],
    ['tags', 'VARCHAR', '255', 'N', 'NULL', '標籤 (逗號分隔)'],
    ['activity_id', 'BIGINT U', '-', 'N', 'NULL', '關聯活動 ID (FK→activities)'],
    ['evidence_url', 'VARCHAR', '255', 'N', 'NULL', '佐證檔案 URL'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
], indexes=[
    'INDEX idx_portfolio_user (user_id)',
])

# ---- TABLE 15: competency_scores ----
add_schema_table(doc, '表 15：職能分數', 'competency_scores', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['user_id', 'BIGINT U', '-', 'Y', '-', '使用者 ID (FK→users, CASCADE)'],
    ['dimension', 'VARCHAR', '50', 'Y', '-', '維度 (leadership/creativity/teamwork/communication/digital)'],
    ['score', 'INT', '-', 'Y', '0', '分數 (0-100)'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (user_id, dimension)',
    'FOREIGN KEY (user_id) → users(id) ON DELETE CASCADE',
    '五維度雷達圖：領導力、創意思維、團隊合作、溝通能力、數位能力',
])

# ---- TABLE 16: ai_review_logs ----
add_schema_table(doc, '表 16：AI 審查記錄', 'ai_review_logs', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['target_type', 'VARCHAR', '50', 'Y', '-', '審查對象類型 (activity/reservation/appeal)'],
    ['target_id', 'BIGINT U', '-', 'Y', '-', '審查對象 ID'],
    ['reviewer_type', 'VARCHAR', '50', 'Y', 'dify_rag', '審查引擎 (dify_rag/gpt4/rule_engine)'],
    ['input_data', 'TEXT', '-', 'N', 'NULL', '輸入資料 (JSON)'],
    ['output_data', 'TEXT', '-', 'N', 'NULL', '輸出結果 (JSON)'],
    ['allow_next_step', 'BOOLEAN', '-', 'Y', 'true', '是否允許進入下一階段'],
    ['risk_level', 'VARCHAR', '20', 'Y', 'Low', "風險等級：Low/Medium/High"],
    ['violations', 'TEXT', '-', 'N', 'NULL', '違規項目列表 (JSON)'],
    ['references', 'TEXT', '-', 'N', 'NULL', '參考法規條文 (JSON)'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
])

# ---- TABLE 17: appeals ----
add_schema_table(doc, '表 17：申訴記錄', 'appeals', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['code', 'VARCHAR', '20', 'Y', '-', '申訴編號 (UNIQUE, e.g. AP-042)'],
    ['appellant_id', 'BIGINT U', '-', 'N', 'NULL', '申訴人 ID (FK→users)'],
    ['appeal_type', 'VARCHAR', '50', 'Y', '-', '申訴類型 (場地衝突/器材損壞/信用積分)'],
    ['subject', 'VARCHAR', '300', 'Y', '-', '申訴主旨'],
    ['description', 'TEXT', '-', 'N', 'NULL', '詳細描述'],
    ['status', 'ENUM', '-', 'Y', 'pending', "狀態：pending/processing/resolved/rejected"],
    ['ai_summary', 'TEXT', '-', 'N', 'NULL', 'AI 生成摘要 (情緒分析+建議)'],
    ['resolution', 'TEXT', '-', 'N', 'NULL', '處理結果'],
    ['assigned_to', 'BIGINT U', '-', 'N', 'NULL', '承辦人 ID (FK→users)'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (code)',
    "CHECK status IN ('pending','processing','resolved','rejected')",
])

# ---- TABLE 18: repairs ----
add_schema_table(doc, '表 18：報修記錄', 'repairs', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['code', 'VARCHAR', '20', 'Y', '-', '報修編號 (UNIQUE, e.g. RP-001)'],
    ['target', 'VARCHAR', '200', 'Y', '-', '報修對象 (投影機/冷氣/音響...)'],
    ['description', 'TEXT', '-', 'Y', '-', '問題描述'],
    ['status', 'ENUM', '-', 'Y', 'pending', "狀態：pending/processing/in_progress/completed/cancelled"],
    ['assignee', 'VARCHAR', '100', 'N', 'NULL', '負責人員'],
    ['submitted_by', 'BIGINT U', '-', 'N', 'NULL', '報修者 ID (FK→users)'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    'UNIQUE (code)',
    "CHECK status IN ('pending','processing','in_progress','completed','cancelled')",
])

# ---- TABLE 19: calendar_events ----
add_schema_table(doc, '表 19：行事曆事件', 'calendar_events', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['title', 'VARCHAR', '200', 'Y', '-', '事件標題'],
    ['date', 'DATE', '-', 'Y', '-', '事件日期'],
    ['type', 'VARCHAR', '50', 'N', 'NULL', '類型 (meeting/competition/performance/workshop/sports/service)'],
    ['color', 'VARCHAR', '20', 'Y', '#003153', '顯示顏色'],
    ['description', 'TEXT', '-', 'N', 'NULL', '事件描述'],
    ['venue', 'VARCHAR', '100', 'N', 'NULL', '地點'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
])

# ---- TABLE 20: time_capsules ----
add_schema_table(doc, '表 20：數位時光膠囊', 'time_capsules', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['club_id', 'BIGINT U', '-', 'N', 'NULL', '社團 ID (FK→clubs)'],
    ['created_by', 'BIGINT U', '-', 'N', 'NULL', '建立者 ID (前任社長, FK→users)'],
    ['received_by', 'BIGINT U', '-', 'N', 'NULL', '接收者 ID (現任社長, FK→users)'],
    ['term', 'VARCHAR', '50', 'Y', '-', '學期 (e.g. 114-1)'],
    ['file_manifest', 'TEXT', '-', 'N', 'NULL', '檔案清單 (JSON)'],
    ['r2_storage_key', 'VARCHAR', '255', 'N', 'NULL', 'Cloudflare R2 儲存金鑰'],
    ['gps_watermark', 'VARCHAR', '255', 'N', 'NULL', 'GPS 浮水印座標'],
    ['status', 'ENUM', '-', 'Y', 'sealed', "狀態：sealed/transferred/received"],
    ['sealed_at', 'TIMESTAMP', '-', 'N', 'NULL', '封裝時間'],
    ['transfer_at', 'TIMESTAMP', '-', 'N', 'NULL', '移交時間'],
    ['received_at', 'TIMESTAMP', '-', 'N', 'NULL', '接收時間'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK status IN ('sealed','transferred','received')",
    '解封需現任社長 + 前任社長雙重驗證',
])

# ---- TABLE 21: outbox ----
add_schema_table(doc, '表 21：事件發送佇列', 'outbox', [
    ['id', 'BIGINT', 'AUTO', 'Y', '-', '主鍵'],
    ['event_type', 'VARCHAR', '100', 'Y', '-', '事件類型 (email/pdf/notification)'],
    ['payload', 'TEXT', '-', 'Y', '-', '事件資料 (JSON)'],
    ['status', 'ENUM', '-', 'Y', 'pending', "狀態：pending/processed/failed"],
    ['retry_count', 'INT', '-', 'Y', '0', '重試次數'],
    ['created_at', 'TIMESTAMP', '-', 'N', 'NULL', '建立時間'],
    ['updated_at', 'TIMESTAMP', '-', 'N', 'NULL', '更新時間'],
], constraints=[
    "CHECK status IN ('pending','processed','failed')",
    'Event Sourcing 模式，確保異步任務可靠執行',
], indexes=[
    'INDEX idx_outbox_status (status)',
])

# ---- 3.4 Key Constraints Summary ----
doc.add_heading('3.4 關鍵約束與索引彙總', level=2)
constraints_summary = [
    'users.role：CHECK IN (student, officer, professor, admin, it)',
    'reservations.priority_level：CHECK IN (1, 2, 3)，Level 1 為校方最高優先',
    'reservations.stage：CHECK IN (algorithm, negotiation, approval, approved, rejected)',
    'credit_logs.action：CHECK IN (add, deduct)，低於 60 分觸發 JWT 失效',
    'venues.status：CHECK IN (available, occupied, maintenance, reserved)',
    '唯一約束：users.student_id, users.email, equipment.code, certificates.certificate_code, appeals.code, repairs.code',
    '外鍵 CASCADE：club_members, activity_registrations, equipment_borrowings, credit_logs, notification_logs, certificates, portfolio_entries, competency_scores',
]
for c in constraints_summary:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('3.5 事務控制', level=2)
doc.add_paragraph(
    '場地預約採用 Pessimistic Locking (SELECT ... FOR UPDATE) 確保並行預約時無 Race Condition。'
    '信用積分扣抵使用 START TRANSACTION 確保一致性。'
    'Outbox 表採 Event Sourcing 模式，確保郵件/PDF/通知等異步任務可靠執行。'
)

# ============================================================
# 4. FRONTEND ARCHITECTURE
# ============================================================
doc.add_heading('4. 前端架構', level=1)

doc.add_heading('4.1 頁面結構 (21 頁)', level=2)
pages_tbl = doc.add_table(rows=1, cols=3)
pages_tbl.style = 'Table Grid'
for i, h in enumerate(['路由', '頁面標題', '說明']):
    cell = pages_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(10)

pages = [
    ['/', 'Landing Page', 'Hero slider / 功能特色 / 十大模組 / 評價 / 公告'],
    ['/login', 'Login', 'Google OAuth + 帳密 + 5 角色 Demo'],
    ['/dashboard?role=X', 'Dashboard', '最近活動 / Leaflet 地圖 / Chart.js 圖表'],
    ['/campus-map', 'Campus Map', 'GeoJSON 建築 + 4 LayerGroups + 搜尋'],
    ['/module/venue-booking', '場地預約', '三階段調度 + 場地 CRUD + 衝突偵測'],
    ['/module/equipment', '設備借用', '借還流程 + CRUD + LINE 通知'],
    ['/module/calendar', '行事曆', 'evo-calendar + 事件 CRUD'],
    ['/module/club-info', '社團資訊', '102 社團 + 搜尋篩選 + CRUD'],
    ['/module/activity-wall', '活動牆', '卡片 + 分類 + 報名 + CRUD'],
    ['/module/ai-overview', 'AI 概覽', 'AI 模組狀態 Dashboard'],
    ['/module/ai-guide', 'AI 企劃生成', 'Dify Workflow 企劃書生成'],
    ['/module/rag-search', 'RAG 法規查詢', 'AI 預審模擬'],
    ['/module/rag-rental', '租借 RAG', '借用規則查詢'],
    ['/module/repair', '報修管理', 'Full CRUD + 追蹤碼'],
    ['/module/appeal', '申訴記錄', 'Full CRUD + AI 摘要'],
    ['/module/reports', '統計報表', 'Chart.js 多維度統計'],
    ['/module/e-portfolio', 'E-Portfolio', '職能成長作品集'],
    ['/module/certificate', '證書自動化', '數位簽章 + 驗證碼'],
    ['/module/time-capsule', '時光膠囊', 'R2 Storage 移交系統'],
    ['/module/2fa', '2FA', 'TOTP + SMS 雙因子驗證'],
]
for pg in pages:
    row = pages_tbl.add_row()
    for i, val in enumerate(pg):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(9)

doc.add_heading('4.2 共用元件 (Layout System)', level=2)
components = [
    'layouts/app.blade.php：全局 HTML 骨架，Tailwind config、CDN 引用、全局 CSS',
    'layouts/shell.blade.php：登入後頁面框架 — 雙層 Header + Sidebar + Main + Footer + Chatbot',
    'Double Header：上白 (校名/連結/角色切換) + 下深藍 (導覽/搜尋/通知/用戶)',
    'Sidebar：四大區塊導覽，active 黃色標記，信用積分儀表板，系統狀態指示',
    'partials/footer.blade.php：深藍背景，校址、社群連結、聯絡資訊、版權聲明',
    'partials/chatbot.blade.php：柴犬輔寶 AI 助理 FAB + 對話面板',
]
for c in components:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('4.3 互動校園地圖（Leaflet.js）', level=2)
doc.add_paragraph(
    '地圖以淨心堂 (25.0359, 121.4323) 為中心，使用 OpenStreetMap 底圖 + GeoJSON 建築分區疊加。'
    '建築 Popup 包含：學院分色標頭、樓層數、狀態（綠/紅）、「我要預約」按鈕。'
    '四大 LayerGroup：教學大樓、無障礙設施、生活服務、交通設施。'
)

# ============================================================
# 5. BACKEND ARCHITECTURE
# ============================================================
doc.add_heading('5. 後端架構', level=1)

doc.add_heading('5.1 API 端點清單 (38+)', level=2)
api_tbl = doc.add_table(rows=1, cols=4)
api_tbl.style = 'Table Grid'
for i, h in enumerate(['HTTP', '端點', 'Controller', '說明']):
    cell = api_tbl.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '003153')
    for p2 in cell.paragraphs:
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r2 in p2.runs:
            r2.font.color.rgb = RGBColor(255,255,255)
            r2.font.bold = True
            r2.font.size = Pt(9)

apis = [
    ['GET','api/users','UserController@index','使用者列表'],
    ['POST','api/users','UserController@store','建立使用者'],
    ['GET','api/users/{id}','UserController@show','查看使用者'],
    ['PUT','api/users/{id}','UserController@update','更新使用者'],
    ['DELETE','api/users/{id}','UserController@destroy','刪除使用者'],
    ['GET','api/clubs','ClubController@index','社團列表 (含篩選/搜尋)'],
    ['POST','api/clubs','ClubController@store','建立社團'],
    ['PUT','api/clubs/{id}','ClubController@update','更新社團'],
    ['DELETE','api/clubs/{id}','ClubController@destroy','刪除社團'],
    ['GET','api/clubs/stats','ClubController@stats','社團統計'],
    ['GET','api/venues','VenueController@index','場地列表'],
    ['POST','api/venues','VenueController@store','建立場地'],
    ['PUT','api/venues/{id}','VenueController@update','更新場地'],
    ['DELETE','api/venues/{id}','VenueController@destroy','刪除場地'],
    ['GET','api/venues/{id}/schedule','VenueController@schedule','場地時段表'],
    ['GET','api/activities','ActivityController@index','活動列表'],
    ['POST','api/activities','ActivityController@store','建立活動 + AI 預審'],
    ['PUT','api/activities/{id}','ActivityController@update','更新活動'],
    ['DELETE','api/activities/{id}','ActivityController@destroy','刪除活動'],
    ['POST','api/activities/{id}/register','ActivityController@register','報名活動'],
    ['POST','api/activities/{id}/cancel-registration','ActivityController@cancelRegistration','取消報名'],
    ['GET','api/equipment','EquipmentController@index','器材列表'],
    ['POST','api/equipment','EquipmentController@store','建立器材'],
    ['PUT','api/equipment/{id}','EquipmentController@update','更新器材'],
    ['DELETE','api/equipment/{id}','EquipmentController@destroy','刪除器材'],
    ['POST','api/equipment/borrow','EquipmentController@borrow','借用器材'],
    ['POST','api/equipment/return','EquipmentController@returnEquipment','歸還器材'],
    ['POST','api/equipment/remind','EquipmentController@remind','LINE 催還提醒'],
    ['GET','api/reservations','ReservationController@index','預約列表'],
    ['POST','api/reservations','ReservationController@store','建立預約 (含衝突偵測)'],
    ['PUT','api/reservations/{id}','ReservationController@update','更新預約'],
    ['DELETE','api/reservations/{id}','ReservationController@destroy','刪除預約'],
    ['POST','api/reservations/{id}/negotiate','ReservationController@negotiate','AI 協商建議'],
    ['POST','api/reservations/{id}/accept-suggestion','ReservationController@acceptSuggestion','接受協商方案'],
    ['GET','api/conflicts','CrudController@conflictIndex','衝突列表'],
    ['POST','api/conflicts','CrudController@conflictStore','建立衝突'],
    ['POST','api/conflicts/negotiate','CrudController@conflictNegotiate','AI 衝突協商'],
    ['GET','api/repairs','CrudController@repairIndex','報修列表'],
    ['POST','api/repairs','CrudController@repairStore','建立報修'],
    ['PUT','api/repairs/{id}','CrudController@repairUpdate','更新報修'],
    ['DELETE','api/repairs/{id}','CrudController@repairDestroy','刪除報修'],
    ['GET','api/appeals','CrudController@appealIndex','申訴列表'],
    ['POST','api/appeals','CrudController@appealStore','建立申訴'],
    ['PUT','api/appeals/{id}','CrudController@appealUpdate','更新申訴'],
    ['DELETE','api/appeals/{id}','CrudController@appealDestroy','刪除申訴'],
    ['POST','api/appeals/{id}/ai-summary','CrudController@appealAiSummary','AI 申訴摘要'],
    ['GET','api/calendar/events','CrudController@calendarIndex','行事曆列表'],
    ['POST','api/calendar/events','CrudController@calendarStore','新增行事曆事件'],
    ['PUT','api/calendar/events/{id}','CrudController@calendarUpdate','更新事件'],
    ['DELETE','api/calendar/events/{id}','CrudController@calendarDestroy','刪除事件'],
    ['GET','api/credits/{userId}','CrudController@creditShow','信用積分查詢'],
    ['POST','api/credits/deduct','CrudController@creditDeduct','扣除信用積分'],
    ['GET','api/notifications/{userId}','CrudController@notificationIndex','通知列表'],
    ['POST','api/notifications/{id}/read','CrudController@notificationRead','標記已讀'],
    ['POST','api/certificates/generate','CrudController@certificateGenerate','產生數位證書'],
    ['GET','api/dashboard/stats/{role}','CrudController@dashboardStats','角色統計數據'],
    ['POST','api/ai/pre-review','CrudController@aiPreReview','AI RAG 預審'],
    ['POST','api/ai/generate-proposal','CrudController@aiGenerateProposal','AI 企劃書生成'],
    ['GET','api/i18n/{lang}','CrudController@i18n','多語系翻譯包'],
    ['GET','api/health','CrudController@health','系統健康檢查'],
]
for a in apis:
    row = api_tbl.add_row()
    for i, val in enumerate(a):
        row.cells[i].text = val
        for p2 in row.cells[i].paragraphs:
            for r2 in p2.runs:
                r2.font.size = Pt(8)

doc.add_heading('5.2 三階段資源調度', level=2)
stages = [
    '第一階段 - 志願序配對（Algorithm）：依 Level 1/2/3 優先權自動配對，無衝突直接進入核准',
    '第二階段 - 自主協商（Negotiation）：偵測衝突後啟動 3/6 分鐘計時器。3 分鐘 GPT-4 介入建議，6 分鐘強制關閉並扣雙方各 10 信用分',
    '第三階段 - 官方核准（Approval）：RAG 法規比對 + Gatekeeping，產出核准文號',
]
for s in stages:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('5.3 AI 整合（Dify + RAG）', level=2)
ai_features = [
    'AI 預審：Dify RAG Engine 比對 5+ 法規文件（Pinecone 向量 DB），返回 risk_level / violations / references / suggestions',
    'AI 企劃生成：Dify Workflow 參考法規模板，產出完整企劃書含預算/流程/SDG 對應',
    'AI 協商監控：GPT-4 分析雙方歷史使用紀錄，3 分鐘後自動提供替代方案',
    'AI 申訴摘要：即時推送案件分析，含情緒分析/緊急度/建議方案',
    'AI 導覽 (輔寶)：柴犬化身 FAB，支援法規查詢/場地導覽/流程引導',
]
for f in ai_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.4 安全機制', level=2)
security = [
    'Cloudflare WAF：防止 SQL Injection / XSS / DDoS',
    'Google OAuth 2.0：SSO 單一登入',
    'TOTP + SMS 2FA：敏感操作雙因子驗證',
    '信用積分制度：低於 60 分強制登出 (Observer Pattern)',
    'CSRF Protection：Laravel 自動 CSRF Token',
    'Rate Limiting：API 請求頻率限制',
]
for s in security:
    doc.add_paragraph(s, style='List Bullet')

# ============================================================
# 6-10 REMAINING SECTIONS (condensed)
# ============================================================
doc.add_heading('6. 核心模組規格', level=1)
modules = [
    ('6.1 職能 E-Portfolio', '記錄學生參與社團活動的職能成長歷程，提供雷達圖分析與 PDF 履歷匯出。', 'portfolio_entries, competency_scores'),
    ('6.2 AI 企劃生成器', '調用 Dify Workflow 參考校內法規，自動生成完整活動企劃書。', 'ai_review_logs'),
    ('6.3 幹部證書自動化', '自動產生社團幹部數位證書，含 SHA256 數位簽章與 QR Code 驗證。', 'certificates'),
    ('6.4 器材盤點追蹤', '完整器材生命週期管理，含借還流程、LINE/SMS 到期提醒、信用積分連動。', 'equipment, equipment_borrowings, credit_logs'),
    ('6.5 AI 智慧預審 (RAG)', 'Dify RAG Engine 自動比對校內法規，標記合規風險，提供修改建議。', 'ai_review_logs'),
    ('6.6 場地活化大數據', '場地使用率熱力圖、周轉排行榜，搭配三階段資源調度系統。', 'venues, reservations, conflicts'),
    ('6.7 AI 申訴摘要', '即時生成申訴案件摘要，含情緒分析與建議方案。', 'appeals'),
    ('6.8 動態活動牆', '卡片式活動展示，支援分類篩選、關鍵字搜尋、即時報名。', 'activities, activity_registrations'),
    ('6.9 數位時光膠囊', '社團移交文件封裝至 R2 Storage，確保跨屆資產安全傳承。', 'time_capsules'),
    ('6.10 全方位 2FA', 'TOTP + SMS 雙因子驗證，強制用於仲裁、核銷、後台登入。', 'users'),
]
for title, desc, tables in modules:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    r = p.add_run(f'相關資料表：{tables}')
    r.font.size = Pt(10)
    r.font.italic = True

doc.add_heading('7. 無障礙校園地圖規格', level=1)
doc.add_heading('7.1 技術實現', level=2)
map_features = [
    'Leaflet.js v1.9.4 + OpenStreetMap 底圖',
    'GeoJSON 建築分區：campus.geojson (120KB+)，學院分色多邊形',
    '座標系統：WGS84 經緯度，地圖中心 (25.0360, 121.4320) - 淨心堂',
    '四大 LayerGroup：教學大樓、無障礙設施、生活服務、交通設施',
    '建築 Popup：學院色彩標頭 + 狀態 + 「我要預約」按鈕',
]
for m in map_features:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('8. 導覽結構', level=1)
doc.add_heading('8.1 Sidebar 四大區塊', level=2)
sidebar_sections = [
    '主要功能：Dashboard / 校園地圖 / 場地預約 / 設備借用 / 行事曆',
    '社群與社團：社團資訊 / 活動牆',
    'AI 核心專區：AI 概覽 / AI 導覽 / RAG 法規查詢',
    '管理與報表：報修管理 / 申訴記錄 / 統計報表',
]
for s in sidebar_sections:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('8.2 信用積分儀表板', level=2)
doc.add_paragraph('Sidebar 頂部顯示信用積分：起始 100 分、進度條、閾值 60 分警告。低於 60 分觸發 JWT 自動失效（Observer Pattern）。')

doc.add_heading('9. 部署與測試', level=1)
doc.add_heading('9.1 Demo 環境', level=2)
doc.add_paragraph('Live Demo：https://3000-ines7d5od0umg4mb1ae8b-b9b802c4.sandbox.novita.ai')
doc.add_paragraph('GitHub：https://github.com/KY0126/SA-v3-')
doc.add_paragraph('技術棧：PHP 8.2 + Laravel 12.56.0 + SQLite (dev) / MySQL 8.0 (prod)')

doc.add_heading('9.2 測試涵蓋', level=2)
tests = [
    '21 頁面路由完整性驗證（所有返回 HTTP 200）',
    '38+ RESTful API 端點回應驗證（GET/POST/PUT/DELETE 全通過）',
    '全部資料實體 Full CRUD 操作測試',
    '角色切換驗證（admin/officer/professor/student/it）',
    '5 國語系切換（zh-TW/zh-CN/en/ja/ko）',
    '地圖 GeoJSON 載入、4 LayerGroup 切換、建築搜尋/篩選',
    '前端零 JavaScript 錯誤（Playwright 驗證）',
]
for t in tests:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('10. 附錄', level=1)
doc.add_heading('10.1 社團完整資料 (114 學年度)', level=2)
doc.add_paragraph('系統內建 114 學年度完整社團資料：社團 66 個 + 學會 36 個 = 共 102 個組織。包含名稱、分類 (學藝/康樂/體育/藝術/服務/宗教/聯誼/學會)、描述、成員數等完整資訊。')

doc.add_heading('10.2 參考法規文件（RAG 知識庫）', level=2)
regs = ['活動申請辦法 v2.1', '場地使用管理規則 v3.0', '經費補助要點 v1.5', '器材借用辦法 v2.0', '安全管理規範 v1.8']
for r in regs:
    doc.add_paragraph(r, style='List Bullet')

# ============================================================
# SAVE
# ============================================================
output_path = '/home/user/webapp/FJU_Smart_Hub_SA_Specification.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
