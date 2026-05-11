import sys
import json
from docx import Document


def append_to_cell(cell, text):
    if not text:
        return
    cell.paragraphs[-1].add_run(str(text))


def set_cell(cell, text):
    if not text:
        return
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    cell.paragraphs[0].add_run(str(text))


def fill_yellow_form(data, template_path, output_path):
    doc = Document(template_path)
    table = doc.tables[0]

    unit_code           = data.get('unit_code', '')
    unit_name           = data.get('unit_name', '')
    activity_name       = data.get('activity_name', '')
    expected_participants = data.get('expected_participants', '')
    staff_count         = data.get('staff_count', '')
    event_date          = data.get('event_date', '')
    start_time          = data.get('start_time', '')
    end_time            = data.get('end_time', '')
    venue_description   = data.get('venue_description', '')
    responsible_person  = data.get('responsible_person', '')
    department          = data.get('department', '')
    contact_phone       = data.get('contact_phone', '')

    time_str = ''
    if event_date:
        time_str = str(event_date)
        if start_time and end_time:
            time_str += f'  {start_time}－{end_time}'

    # Row 0: ※單位代碼 (col 0) — append value after label text
    append_to_cell(table.cell(0, 0), unit_code)

    # Row 4: 活動負責人 name (col 6) — replace placeholder
    if responsible_person:
        set_cell(table.cell(4, 6), responsible_person)

    # Row 4: 系級 (col 12) — append after "系級："
    append_to_cell(table.cell(4, 12), department)

    # Row 4: 聯絡電話 (col 17) — append after "聯絡電話："
    append_to_cell(table.cell(4, 17), contact_phone)

    # Row 6: 單位名稱 value (col 4) — empty cell
    set_cell(table.cell(6, 4), unit_name)

    # Row 7: 活動名稱 value (col 4) — empty cell
    set_cell(table.cell(7, 4), activity_name)

    # Row 8: 活動對象人數 value (col 4) — empty cell
    set_cell(table.cell(8, 4), expected_participants)

    # Row 8: 工作人員人數 (col 11) — no separate value cell, append to label with separator
    if staff_count:
        append_to_cell(table.cell(8, 11), '：' + str(staff_count))

    # Row 9: 活動起訖時間 value (col 4) — empty cell
    set_cell(table.cell(9, 4), time_str)

    # Row 11: 課指組所轄場地 — 場地名稱 (col 4) — empty cell
    set_cell(table.cell(11, 4), venue_description)

    doc.save(output_path)


if __name__ == '__main__':
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)
    template_path = data.pop('template_path')
    output_path = data.pop('output_path')
    fill_yellow_form(data, template_path, output_path)
    print('OK')
